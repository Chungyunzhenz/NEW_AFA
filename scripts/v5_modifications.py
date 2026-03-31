"""v4→v5: 10 項修改 + 參考文獻補充"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = 'FF0000'
INPUT = '農糧署研究計畫書_v4_0330_v3.docx'
OUTPUT = '農糧署研究計畫書_v5_0330.docx'

# ============================================================
# Helpers
# ============================================================
def mke(text, bold=False):
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), RED); rPr.append(c)
    if bold:
        rPr.append(OxmlElement('w:b'))
    run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; t.set(qn('xml:space'), 'preserve')
    run.append(t)
    return run

def sp(para, text, bold=False):
    """Set paragraph to red text."""
    for ch in list(para._element):
        if (ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag) == 'r':
            para._element.remove(ch)
    if text:
        para._element.append(mke(text, bold))

def cp(para):
    """Clear paragraph."""
    sp(para, '')

def ip(ref, text, bold=False, style_id=None):
    """Insert paragraph after ref element."""
    p = OxmlElement('w:p')
    if style_id:
        pPr = OxmlElement('w:pPr')
        s = OxmlElement('w:pStyle'); s.set(qn('w:val'), style_id)
        pPr.append(s); p.append(pPr)
    if text:
        p.append(mke(text, bold))
    ref.addnext(p)
    return p

def it(ref, headers, rows):
    """Insert table after ref element."""
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    ts = OxmlElement('w:tblStyle'); ts.set(qn('w:val'), 'TableGrid'); tblPr.append(ts)
    tw = OxmlElement('w:tblW'); tw.set(qn('w:w'), '0'); tw.set(qn('w:type'), 'auto'); tblPr.append(tw)
    borders = OxmlElement('w:tblBorders')
    for bn in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'0'); b.set(qn('w:color'),'000000')
        borders.append(b)
    tblPr.append(borders); tbl.append(tblPr)
    g = OxmlElement('w:tblGrid')
    for _ in headers: g.append(OxmlElement('w:gridCol'))
    tbl.append(g)
    def mc(t, hd=False):
        tc = OxmlElement('w:tc'); p = OxmlElement('w:p'); p.append(mke(str(t), hd)); tc.append(p); return tc
    tr = OxmlElement('w:tr')
    for h in headers: tr.append(mc(h, True))
    tbl.append(tr)
    for rd in rows:
        tr = OxmlElement('w:tr')
        for ct in rd: tr.append(mc(ct))
        tbl.append(tr)
    ref.addnext(tbl)
    return tbl

def sc(cell, text, bold=False):
    """Set cell text red."""
    for p in cell.paragraphs:
        for ch in list(p._element):
            if (ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag) == 'r':
                p._element.remove(ch)
    cell.paragraphs[0]._element.append(mke(text, bold))

# ============================================================
doc = Document(INPUT)
P = doc.paragraphs
T = doc.tables

# ============================================================
# 1. 1.3 移除詳細資料筆數
# ============================================================
print("1. 1.3 移除資料筆數...")
sp(P[14], '本計畫之預測模型建立於多源政府開放資料之整合基礎上，涵蓋多種作物、批發市場及縣市。')
sp(P[15], '資料來源包括：農糧署批發市場交易行情、中央氣象署氣象觀測、農情產量統計及颱風資料庫，'
   '詳見表 1-2。整合過程面臨更新頻率不一、格式標準不同、跨機關缺乏統一串接機制等困境，'
   '本計畫透過結構化資料庫與自動化 API 擷取管道加以解決。')

# ============================================================
# 2. 1.4 三模型表格
# ============================================================
print("2. 1.4 插入三模型表格...")
sp(P[26],
   '綜合觀之，目前臺灣學術文獻中尚無研究同時達成以下三項條件：採用多模型機器學習方法、'
   '整合氣象與產銷多源開放資料、針對甘藍進行價格預測。此即本計畫之學術切入點。'
   '本研究採用以下三種獨立預測模型：')
tbl_model = it(P[26]._element,
    ["模型", "技術類型", "Python 模組", "核心優勢"],
    [["Prophet", "可分解式時間序列",
      "prophet (Prophet)",
      "自動偵測趨勢與季節性，適合年度週期明顯之農產品（Taylor & Letham, 2018）"],
     ["XGBoost", "梯度提升樹",
      "xgboost (XGBRegressor)",
      "特徵交互學習力強，提供特徵重要性排名，可解釋性高（Chen & Guestrin, 2016）"],
     ["ANN", "人工神經網路",
      "sklearn.neural_network\n(MLPRegressor)",
      "學習複雜非線性映射，與樹模型互補（Rumelhart et al., 1986）"]])
ip(tbl_model,
   '三種模型分別代表可分解式、梯度提升樹與神經網路三種不同技術路線，'
   '透過多模型比較分析，探討各方法在不同蔬菜品項與市場情境下之預測表現差異。')

# ============================================================
# 3. 第二章：只保留 XGBoost、ANN、Prophet
# ============================================================
print("3. 重構第二章文獻回顧...")

sp(P[49], '本節就本研究所採用之三種核心預測方法——XGBoost、Prophet 與 ANN——'
   '歸納其理論依據、技術特性與農業預測情境下之適用性，作為後續研究設計之理論基礎。')

# 2.1.1 → XGBoost
sp(P[50], '2.1.1　梯度提升樹方法（XGBoost）')
sp(P[51],
   'Chen 與 Guestrin（2016）提出之 XGBoost（eXtreme Gradient Boosting），'
   '透過正則化目標函數與高效之分裂演算法，建立數百棵決策樹逐步修正殘差，'
   '在各類預測競賽中取得優異成績。其核心優勢在於特徵工程之靈活性——'
   '研究者可將滯後值、滾動統計量、日曆特徵、氣象變數等多種資訊納入模型，'
   '學習複雜的非線性交互效應。此外，XGBoost 提供特徵重要性排名，'
   '使預測結果具備高度可解釋性，能回答「哪些因素影響了價格預測」。')
sp(P[52],
   '然而，樹模型之決策邊界為分段常數函數，對時間序列之長期趨勢與季節性成分建模不如專門的時間序列模型直觀，'
   '且對平滑非線性關係之捕獲不如神經網路方法。')

# 2.1.2 → Prophet
sp(P[53], '2.1.2　可分解式預測模型（Prophet）')
sp(P[54],
   'Taylor 與 Letham（2018）發表之 Prophet 模型，採用可加或可乘式分解架構，'
   '將時間序列拆解為趨勢、季節性、假日效應與外生回歸項等成分。'
   'Prophet 之設計哲學為「讓分析師更容易產出合理的預測」，'
   '其自動化之變化點偵測（changepoint detection）與彈性之季節性建模，'
   '特別適合處理具明顯年度週期性之農產品資料，且可納入天氣、颱風等外生變數作為額外回歸項。')
sp(P[55],
   'Prophet 之限制在於其本質上仍是曲線配適方法，對短期序列相關性之捕獲能力較弱，'
   '且在資料量較少時預測穩定性較低。因此需搭配其他類型之模型互補。')

# 2.1.3 → ANN
sp(P[56], '2.1.3　人工神經網路（ANN）')
sp(P[57],
   '人工神經網路（Artificial Neural Network）之理論基礎源於 Rumelhart、Hinton 與 Williams（1986）'
   '提出之反向傳播演算法（Backpropagation），透過多層感知器（MLP）架構，'
   '可學習輸入特徵與目標變數之間的複雜非線性映射關係。'
   '本研究採用 Scikit-learn 套件（Pedregosa et al., 2011）提供之 MLPRegressor，'
   '其參數量適中，在約 170 個月度資料點之規模下較不易過擬合。'
   'ANN 與 XGBoost 之關鍵差異在於：XGBoost 產生分段常數之決策函數，'
   '而 ANN 透過連續可微之激活函數產生平滑之函數逼近，兩者具有互補性。')
sp(P[58],
   'ANN 之限制在於可解釋性較低，模型內部權重難以直觀解讀，'
   '但可透過排列重要性（permutation importance）等事後分析方法補強。'
   '此外，ANN 之超參數調校（隱藏層數、神經元數、學習率等）對預測表現影響較大，'
   '需透過系統性實驗確定最佳配置。')

# 2.1.4 原 ANN → 清除（已移至 2.1.3）
sp(P[59], '2.1.4　三種預測方法綜合比較')
cp(P[60])

# 2.1.5 → 清除
cp(P[62])

# 更新方法比較表 (找到正確的表格)
print("   更新方法比較表...")
for ti, table in enumerate(T):
    if table.cell(0,0).text.strip() in ('方法類別', '梯度提升樹'):
        # This is the comparison table - rebuild it
        # Row 0: headers (keep)
        # Row 1: was 傳統統計 → XGBoost
        sc(table.cell(1,0), "梯度提升樹")
        sc(table.cell(1,1), "XGBoost\n(Chen & Guestrin, 2016)\nPython: xgboost")
        sc(table.cell(1,2), "特徵工程靈活、非線性建模能力強、特徵重要性排名提供高可解釋性")
        sc(table.cell(1,3), "對時間序列趨勢與季節性建模不如專門模型")
        # Row 2: was 機器學習 → Prophet
        sc(table.cell(2,0), "可分解式")
        sc(table.cell(2,1), "Prophet\n(Taylor & Letham, 2018)\nPython: prophet")
        sc(table.cell(2,2), "自動變化點偵測、季節性建模彈性高、可納入外生變數")
        sc(table.cell(2,3), "本質為曲線配適，短期序列相關性捕獲較弱")
        # Row 3: was 可分解式 → ANN
        sc(table.cell(3,0), "人工神經網路")
        sc(table.cell(3,1), "ANN / MLP\n(Rumelhart et al., 1986)\nPython: sklearn MLPRegressor")
        sc(table.cell(3,2), "學習複雜非線性映射、與梯度提升樹互補、中小規模資料適用")
        sc(table.cell(3,3), "可解釋性較低、超參數調校敏感")
        # Row 4: was 集成/ANN → 清除（刪除行比較複雜，改為標記刪除）
        if len(table.rows) > 4:
            sc(table.cell(4,0), "—"); sc(table.cell(4,1), "—"); sc(table.cell(4,2), "—"); sc(table.cell(4,3), "—")
        print(f"   已更新 Table {ti}")
        break

# ============================================================
# 4. 第三章核心問題簡化
# ============================================================
print("4. 簡化核心問題...")
sp(P[78],
   '臺灣蔬菜產銷長期存在價格劇烈波動之困境，農民與政策端均缺乏有效之前瞻預測工具。'
   '本研究欲解決以下核心問題：')
sp(P[80],
   '問題一：農民缺乏未來價格趨勢之參考依據，難以做出合理之種植決策，需要一套實用的價格預測方法。')
cp(P[81]) # empty after
sp(P[82],
   '問題二：現有平臺僅提供歷史查詢，缺乏整合氣象與颱風因素之預測分析功能，'
   '需建立能納入外生變數之預測模型。')
cp(P[83])
sp(P[84],
   '問題三：學術研究之預測模型多停留在論文階段，'
   '需將模型成果轉化為可實際操作之決策支援工具。')

# ============================================================
# 5. 刪除 3.1.2 服務對象
# ============================================================
print("5. 刪除 3.1.2 服務對象...")
cp(P[86])  # heading
cp(P[87])  # content
# Also need to clear the service target table
for ti, table in enumerate(T):
    if len(table.rows) > 0 and '服務對象' in table.cell(0,0).text:
        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                sc(table.cell(r,c), '')
        print(f"   清除服務對象表格 Table {ti}")
        break

# ============================================================
# 6. RQ1-RQ3 調整
# ============================================================
print("6. 調整 RQ...")
sp(P[93], '基於上述核心問題，本研究提出以下研究問題：')

sp(P[95], 'RQ1：如何整合來自不同政府機關之農業、氣象與颱風資料，建立適合預測模型使用之結構化資料庫？', bold=True)
sp(P[96], '子問題 1a：不同資料源之時間格式與更新頻率差異如何處理，以產出一致之月度分析資料？')
sp(P[97], '子問題 1b：氣象觀測站與行政區域之空間對應關係如何建立？')
cp(P[98])  # 移除 1c 資料品質燈號（太細節）

sp(P[100], 'RQ2：Prophet、XGBoost 與 ANN 三種模型應用於臺灣蔬菜價格預測之表現各為何？各有何適用情境？', bold=True)
sp(P[101], '子問題 2a：三種模型在不同蔬菜品項上之預測精度是否存在顯著差異？')
sp(P[102], '子問題 2b：納入颱風多維特徵是否能有效提升颱風敏感作物之預測表現？')
cp(P[103])  # 移除 2c（與 2b 合併）

sp(P[105], 'RQ3：如何將預測模型之成果轉化為可供實務使用之互動式決策支援工具？', bold=True)
sp(P[106], '子問題 3a：預測結果以何種視覺化方式呈現最能輔助決策？')
sp(P[107], '子問題 3b：如何設計預警機制，使使用者能即時掌握價格異常訊號？')
cp(P[108])  # 移除 3c

# ============================================================
# 7. 3.4 改善呈現
# ============================================================
print("7. 改善 3.4...")
sp(P[111],
   '本研究以甘藍、萵苣、小白菜、花椰菜、青蔥五種蔬菜為核心對象，'
   '採用 Prophet、XGBoost 與 ANN 三種模型進行月度平均價格之 1、3、6 個月預測。'
   '資料來源涵蓋農糧署交易行情、中央氣象署氣象觀測、農情產量統計及颱風資料庫。'
   '地理範圍為臺灣本島及離島（22 縣市、20 個批發市場）。')

# Clear old "研究範圍：" bold heading and replace 限制 section
sp(P[113],
   '研究限制方面：（1）採月度聚合進行預測，無法捕捉日內或週內之短期波動；'
   '（2）受限於月度資料點數量，未採用深度學習模型；'
   '（3）以月均價格為主要預測目標，交易量為輔助參考。')

# Clear the old scope table and limitation table
# Find and clear them (tables after 3.4 heading)
scope_cleared = 0
for ti, table in enumerate(T):
    if len(table.rows) > 0:
        first = table.cell(0,0).text.strip()
        if first in ('維度', '限制項目') and scope_cleared < 2:
            for r in range(len(table.rows)):
                for c in range(len(table.columns)):
                    sc(table.cell(r,c), '')
            scope_cleared += 1
            print(f"   清除 3.4 表格 Table {ti} ({first})")

# ============================================================
# 8. 刪除第四章，第五章改為第四章
# ============================================================
print("8. 刪除第四章，重編第五章...")

# Clear all Ch4 content (paragraphs 116-175)
for i in range(116, 176):
    if i < len(P):
        cp(P[i])

# Clear all Ch4 tables (tables that belong to Ch4)
# Tables 10-21 are Ch4 tables (after Ch3's tables)
for ti, table in enumerate(T):
    if ti >= 10 and ti <= 21:
        for r in range(len(table.rows)):
            for c in range(len(table.columns)):
                try:
                    sc(table.cell(r,c), '')
                except:
                    pass

# Rename Ch5 → Ch4
sp(P[177], '第四章　工作內容重點')
sp(P[178], '4.1　工作項目')
sp(P[181], '4.2　預期貢獻')
sp(P[187], '4.3　關鍵績效指標（KPI）')

# ============================================================
# 9. KPI 保守化
# ============================================================
print("9. KPI 保守化...")
sp(P[188], '本計畫之關鍵績效指標如下：')
# Find KPI table and update
for ti, table in enumerate(T):
    if len(table.rows) > 0 and 'KPI' in table.cell(0,0).text:
        sc(table.cell(1,0), "預測精度")
        sc(table.cell(1,1), "達合理預測水準")
        sc(table.cell(1,2), "三種模型在五種蔬菜之平均 MAPE 與 R² 表現")
        sc(table.cell(2,0), "模型比較")
        sc(table.cell(2,1), "完成三模型之系統性比較分析")
        sc(table.cell(2,2), "各模型在不同作物上之表現差異報告")
        sc(table.cell(3,0), "資料整合")
        sc(table.cell(3,1), "完成多源資料之串接與入庫")
        sc(table.cell(3,2), "結構化資料庫建置完成")
        sc(table.cell(4,0), "系統開發")
        sc(table.cell(4,1), "完成決策支援平臺之開發與展示")
        sc(table.cell(4,2), "前後端系統可正常運行並展示預測結果")
        if len(table.rows) > 5:
            sc(table.cell(5,0), "—"); sc(table.cell(5,1), "—"); sc(table.cell(5,2), "—")
        print(f"   已更新 KPI Table {ti}")
        break

# ============================================================
# 10. 時程改為 9 個月
# ============================================================
print("10. 時程改為 9 個月...")
sp(P[190], '4.4　專案時程規劃（九個月）')
sp(P[191], '本計畫以九個月為執行期程，各階段規劃如下：')
# Find schedule table and rebuild
for ti, table in enumerate(T):
    if len(table.rows) > 0 and '月份' in table.cell(0,0).text:
        # Clear and rebuild
        sc(table.cell(0,0), "階段"); sc(table.cell(0,1), "月份"); sc(table.cell(0,2), "主要工作"); sc(table.cell(0,3), "里程碑")
        sc(table.cell(1,0), "資料整合與特徵工程"); sc(table.cell(1,1), "第 1–3 個月")
        sc(table.cell(1,2), "多源資料串接與入庫、資料品質檢核、特徵工程設計與實作")
        sc(table.cell(1,3), "結構化資料庫與特徵矩陣完成")
        sc(table.cell(2,0), "模型開發與評估"); sc(table.cell(2,1), "第 4–6 個月")
        sc(table.cell(2,2), "Prophet、XGBoost、ANN 三模型訓練、五種蔬菜預測精度評估與比較分析")
        sc(table.cell(2,3), "三模型績效比較報告完成")
        sc(table.cell(3,0), "系統開發與結案"); sc(table.cell(3,1), "第 7–9 個月")
        sc(table.cell(3,2), "前端儀表板開發、系統整合測試、結案報告撰寫")
        sc(table.cell(3,3), "系統展示與結案報告繳交")
        # Clear extra rows
        for r in range(4, len(table.rows)):
            for c in range(len(table.columns)):
                sc(table.cell(r,c), '')
        print(f"   已更新時程 Table {ti}")
        break

# ============================================================
# 11. 參考文獻補充
# ============================================================
print("11. 補充參考文獻...")
# Find the last reference paragraph and add new ones
# Also remove Box & Jenkins, Hyndman (2008), Ke et al. since no longer cited

# Remove references no longer cited
for i in range(193, len(P)):
    text = P[i].text
    if 'Box, G. E. P.' in text:
        cp(P[i]); print("   移除 Box & Jenkins (1970)")
    elif 'Hyndman, R. J., & Khandakar' in text:
        cp(P[i]); print("   移除 Hyndman & Khandakar (2008)")
    elif 'Ke, G., Meng' in text:
        cp(P[i]); print("   移除 Ke et al. (2017)")

# Add new references - insert after last existing reference
last_ref = None
for i in range(len(P)-1, 190, -1):
    if P[i].text.strip() and not P[i].text.startswith('第') and i > 193:
        last_ref = i
        break

if last_ref:
    ref2 = ip(P[last_ref]._element,
        'Rumelhart, D. E., Hinton, G. E., & Williams, R. J. (1986). '
        'Learning representations by back-propagating errors. Nature, 323(6088), 533–536. '
        'https://doi.org/10.1038/323533a0')
    ref3 = ip(ref2,
        'Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., '
        'Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., '
        'Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, É. (2011). '
        'Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825–2830.')
    print("   新增 Rumelhart et al. (1986)")
    print("   新增 Pedregosa et al. (2011)")

# ============================================================
# Save
# ============================================================
print(f"\nSaving to {OUTPUT}...")
doc.save(OUTPUT)
print("Done!")
