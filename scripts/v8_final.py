"""v7→v8: 濃縮第一章、補空白內容、清理"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = 'FF0000'
INPUT = '農糧署研究計畫書_v7_0330.docx'
OUTPUT = '農糧署研究計畫書_v8_0330.docx'

def mke(text, bold=False, color=RED):
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    if bold: rPr.append(OxmlElement('w:b'))
    run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; t.set(qn('xml:space'), 'preserve')
    run.append(t)
    return run

def sp(para, text, bold=False, color=RED):
    for ch in list(para._element):
        if (ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag) == 'r':
            para._element.remove(ch)
    if text: para._element.append(mke(text, bold, color))

def cp(para):
    sp(para, '')

def ip(ref, text, bold=False, color=RED):
    p = OxmlElement('w:p')
    if text: p.append(mke(text, bold, color))
    ref.addnext(p)
    return p

def it(ref, headers, rows):
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    ts = OxmlElement('w:tblStyle'); ts.set(qn('w:val'), 'TableGrid'); tblPr.append(ts)
    tw = OxmlElement('w:tblW'); tw.set(qn('w:w'), '0'); tw.set(qn('w:type'), 'auto'); tblPr.append(tw)
    borders = OxmlElement('w:tblBorders')
    for bn in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:space'),'0'); b.set(qn('w:color'),'000000')
        borders.append(b)
    tblPr.append(borders); tbl.append(tblPr)
    g = OxmlElement('w:tblGrid')
    for _ in headers: g.append(OxmlElement('w:gridCol'))
    tbl.append(g)
    def mc(t, hd=False):
        tc = OxmlElement('w:tc'); p = OxmlElement('w:p'); p.append(mke(str(t), hd)); tc.append(p); return tc
    tr = OxmlElement('w:tr')
    for h in headers: tr.append(mc(h, True))
    tbl.append(tr)
    for rd in rows:
        tr = OxmlElement('w:tr')
        for ct in rd: tr.append(mc(ct))
        tbl.append(tr)
    ref.addnext(tbl)
    return tbl

doc = Document(INPUT)
P = doc.paragraphs

# ============================================================
# A. 濃縮第一章：合併 1.1+1.2 → 1.1，刪除 1.3/1.4 標題
#    前面放計畫亮點
# ============================================================
print("A. 濃縮第一章...")

# [1] 1.1 改為更精簡的標題
sp(P[1], '1.1　研究背景', bold=True)

# [2] 合併 1.1+1.2 核心內容為一段
sp(P[2],
   '臺灣蔬菜產銷長期面臨「菜金菜土」之結構性困境。以甘藍為例，其佔全部蔬菜產量之 17%（焦鈞，2024），'
   '價格走揚時農民搶種致供過於求，颱風後消費者搶購又致價格翻倍。'
   'Su et al.（2025）之實證研究指出，政策介入之「時機」較「手段」更具決定性。'
   '然而，現行供苗預警機制因回報數據不完整、預測精準度不足等問題，'
   '難以在農民做出種植決策前提供可靠之價格趨勢預測（今周刊，2018；報導者，2021）。'
   '現有政府平臺均僅提供歷史查詢，尚無以機器學習進行蔬菜價格前瞻預測之系統。')

# [3] 改為計畫目標一句話
sp(P[3],
   '本計畫之核心目標，即透過 AI 預測系統補強「事前預知」能力，'
   '使農糧署得以在產銷異常發生之前啟動政策工具。')

# 清除原 1.2 所有段落
for i in [5, 6, 7, 8, 9]:
    if i < len(P):
        cp(P[i])
print("  已合併 1.1+1.2")

# [5] 原 1.2 標題 → 改為「1.2 計畫亮點」
sp(P[5], '1.2　計畫亮點', bold=True)

# [6] 插入亮點摘要
sp(P[6], '本計畫之主要特色與創新如下：')

# [7] 用原 P[7] 的位置放亮點表格引導
sp(P[7], '')  # clear

# Insert highlights table after P[6]
tbl_hl = it(P[6]._element,
    ["亮點", "說明"],
    [["三模型獨立比較",
      "採用 Prophet、XGBoost、ANN 三種不同技術路線之預測模型，系統性比較各方法之預測表現與適用情境"],
     ["五種關鍵蔬菜",
      "以甘藍、萵苣、小白菜、花椰菜、青蔥為驗證對象，涵蓋葉菜、花菜、辛香料三類"],
     ["多源資料整合",
      "整合農糧署交易行情、中央氣象署氣象觀測、農情產量統計及颱風資料庫四大政府開放資料"],
     ["颱風特徵工程",
      "設計七維颱風特徵向量（強度、路徑、延遲效應等），超越既有研究僅用二元變數之做法"],
     ["互動式決策支援",
      "開發 Web 平臺，含價格預測圖、臺灣縣市地圖、颱風情境模擬、紅黃綠預警燈號"]])

print("  已插入計畫亮點表格")

# 清除 P[8]-[9]（原 1.2 剩餘段落）
for i in [8, 9]:
    if i < len(P):
        cp(P[i])

# [11] 原 1.3 標題 → 改為 1.3（保持不變但簡化內容）
sp(P[11], '1.3　開放資料來源', bold=True)

# [12]-[13] 保持（已精簡過）

# [15] 平臺比較段落 → 精簡
sp(P[15], '表 1-1 彙整現有平臺與本計畫之功能比較，顯示目前尚無平臺具備以 ML 進行蔬菜價格前瞻預測之能力。')

# [17]-[20] 表格標題和結論 → 清除多餘的
cp(P[17])  # "表 1-1 臺灣主要..." 標題文字（表格自己有）
cp(P[20])  # 結論段落（已在 P[15] 說明）

# [22] 原 1.4 → 改為 1.4（簡化）
sp(P[22], '1.4　學術定位', bold=True)

# [23] Peng 文獻段落 → 精簡
sp(P[23],
   'Peng et al.（2015）為臺灣最早系統性比較多種演算法進行農產品價格預測之研究。'
   '然而，該研究僅使用歷史交易價格，未納入氣象變數，亦未採用多模型比較架構。'
   '目前臺灣學術文獻中尚無研究同時達成：採用多模型 ML 方法、整合氣象與產銷多源開放資料、'
   '針對蔬菜進行價格預測——此即本計畫之學術切入點。')

# [24] 模型介紹 → 精簡（表格已在 1.2 亮點中呈現）
sp(P[24], '本研究採用之三種預測模型詳見下表：')
# 表格 1（模型表）已存在，保持

# [25]-[26] 精簡
sp(P[25], '三種模型分別代表可分解式、梯度提升樹與神經網路三種技術路線，透過多模型比較分析探討各方法之適用情境。')
sp(P[26],
   '本計畫以甘藍為主要分析標的，搭配萵苣、小白菜、花椰菜、青蔥共五種蔬菜作為驗證對象。')

# 清除資料盤點段落的冗餘文字（表格自己會說話）
sp(P[28], '五種蔬菜之資料盤點', bold=True)
sp(P[29], '以下為五種蔬菜之資料量與完整性：')

sp(P[32], '五種蔬菜之選定理由', bold=True)

sp(P[35], '輔助資料來源', bold=True)

# 清除重複的結論段落
cp(P[37])  # "五種蔬菜涵蓋..." 重複

# 清除 "資料來源" 獨立標題段落
for i, p in enumerate(P):
    if p.text.strip() == '資料來源' and i < 45:
        cp(p)
        break

print("  第一章濃縮完成")

# ============================================================
# B. 填補空白內容
# ============================================================
print("\nB. 填補空白內容...")

# 2.1.4 三種預測方法綜合比較 → 加內容
for i, p in enumerate(P):
    if '三種預測方法綜合比較' in p.text:
        # Insert content after heading
        p_content = ip(p._element,
            '下表綜合比較三種方法之核心特性（詳見表 2-1）。'
            '三種模型在歸納偏差上各有不同：Prophet 以趨勢分解見長，XGBoost 以特徵交互見長，'
            'ANN 以非線性函數逼近見長，三者互補可涵蓋農產品價格之多重影響因素。')
        print(f"  P[{i}] 2.1.4 已補充內容")
        break

# 2.4 研究缺口 → 加內容
for i, p in enumerate(P):
    if '本研究識別出以下' in p.text and '研究缺口' in p.text:
        sp(p, '綜合上述文獻回顧，本研究識別出以下研究缺口：')
        # The gap table (Table 7) should be right after this - check if it has content
        print(f"  P[{i}] 2.4 研究缺口已確認")
        break

# Fix "詳見第四章" reference → 第四章已刪除
for i, p in enumerate(P):
    if '詳見第四章' in p.text:
        new_text = p.text.replace('（詳見第四章）', '').replace('詳見第四章', '')
        sp(p, new_text)
        print(f"  P[{i}] 修正「詳見第四章」引用")

# ============================================================
# C. 最終清理空白段落
# ============================================================
print("\nC. 清理空白段落...")
body = doc.element.body
removed = 0
prev_empty = False
for elem in list(body):
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    if tag == 'p':
        text = ''.join(t.text or '' for t in elem.iter() if t.tag.endswith('}t') or t.tag == 't').strip()
        has_heading = False
        for pPr in elem.findall(qn('w:pPr')):
            for pStyle in pPr.findall(qn('w:pStyle')):
                if 'Heading' in pStyle.get(qn('w:val'), ''):
                    has_heading = True
        if not text and not has_heading:
            if prev_empty:
                body.remove(elem); removed += 1
            prev_empty = True
        else:
            prev_empty = False
    else:
        prev_empty = False

# Remove empty headings that were cleared
for elem in list(body):
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    if tag == 'p':
        text = ''.join(t.text or '' for t in elem.iter() if t.tag.endswith('}t') or t.tag == 't').strip()
        if not text:
            for pPr in elem.findall(qn('w:pPr')):
                for pStyle in pPr.findall(qn('w:pStyle')):
                    if 'Heading' in pStyle.get(qn('w:val'), ''):
                        body.remove(elem); removed += 1

print(f"  移除 {removed} 個空白元素")

# ============================================================
# Save
# ============================================================
print(f"\nSaving to {OUTPUT}...")
doc.save(OUTPUT)

# Verify
doc2 = Document(OUTPUT)
print(f"最終段落: {len(doc2.paragraphs)}, 表格: {len(doc2.tables)}")
empty = sum(1 for p in doc2.paragraphs if not p.text.strip())
print(f"剩餘空白段落: {empty}")
print("Done!")
