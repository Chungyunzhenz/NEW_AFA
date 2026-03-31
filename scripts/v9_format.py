"""v8→v9: 全部黑色字體、統一表格樣式、格式調整"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLACK = RGBColor(0, 0, 0)
INPUT = '農糧署研究計畫書_v8_final.docx'
OUTPUT = '農糧署研究計畫書_v9_final.docx'

doc = Document(INPUT)

# ============================================================
# 1. 所有段落：字體統一黑色 + 字型大小
# ============================================================
print("1. 統一段落字體顏色與格式...")
para_count = 0
for p in doc.paragraphs:
    for run in p.runs:
        run.font.color.rgb = BLACK
        if not run.font.size:
            style_name = p.style.name if p.style else ''
            if 'Heading 1' in style_name:
                run.font.size = Pt(16)
                run.bold = True
            elif 'Heading 2' in style_name:
                run.font.size = Pt(14)
                run.bold = True
            elif 'Heading 3' in style_name:
                run.font.size = Pt(13)
                run.bold = True
            else:
                run.font.size = Pt(12)
        para_count += 1

    # Also fix color in XML directly (for runs created via OxmlElement)
    for elem in p._element.iter():
        if elem.tag.endswith('}color') or elem.tag == qn('w:color'):
            elem.set(qn('w:val'), '000000')

print(f"   已處理 {para_count} 個 run")

# ============================================================
# 2. 所有表格：字體黑色 + 統一樣式
# ============================================================
print("\n2. 統一表格樣式...")
for ti, table in enumerate(doc.tables):
    # Set table alignment to center
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table style
    tbl_elem = table._tbl
    tblPr = tbl_elem.find(qn('w:tblPr'))
    if tblPr is None:
        from docx.oxml import OxmlElement
        tblPr = OxmlElement('w:tblPr')
        tbl_elem.insert(0, tblPr)

    # Set table width to 100%
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        from docx.oxml import OxmlElement
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')

    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            # Fix all text color to black
            for p in cell.paragraphs:
                # Set paragraph alignment
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for run in p.runs:
                    run.font.color.rgb = BLACK
                    run.font.size = Pt(10)
                    # Header row bold
                    if ri == 0:
                        run.bold = True
                    else:
                        run.bold = False

                # Fix XML-level color
                for elem in p._element.iter():
                    if elem.tag.endswith('}color') or elem.tag == qn('w:color'):
                        elem.set(qn('w:val'), '000000')

            # Set cell margins for padding
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                from docx.oxml import OxmlElement
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            # Cell margins (top/bottom/left/right)
            tcMar = tcPr.find(qn('w:tcMar'))
            if tcMar is None:
                from docx.oxml import OxmlElement
                tcMar = OxmlElement('w:tcMar')
                tcPr.append(tcMar)
            for side in ['top', 'bottom', 'start', 'end']:
                from docx.oxml import OxmlElement
                margin = tcMar.find(qn(f'w:{side}'))
                if margin is None:
                    margin = OxmlElement(f'w:{side}')
                    tcMar.append(margin)
                margin.set(qn('w:w'), '60')
                margin.set(qn('w:type'), 'dxa')

            # Header row shading (light gray background)
            if ri == 0:
                shading = tcPr.find(qn('w:shd'))
                if shading is None:
                    from docx.oxml import OxmlElement
                    shading = OxmlElement('w:shd')
                    tcPr.append(shading)
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'D9E2F3')  # Light blue-gray

    print(f"   Table {ti}: {len(table.rows)}行 — 已統一")

# ============================================================
# 3. 參考文獻格式調整：移除藍色備註標題
# ============================================================
print("\n3. 調整參考文獻格式...")
for i, p in enumerate(doc.paragraphs):
    if '【以下為新補充之文獻' in p.text:
        # Remove this annotation line
        for ch in list(p._element):
            tag = ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag
            if tag == 'r':
                p._element.remove(ch)
        print(f"   P[{i}] 移除藍色備註標題")

# ============================================================
# 4. 段落間距統一
# ============================================================
print("\n4. 統一段落間距...")
for p in doc.paragraphs:
    style_name = p.style.name if p.style else ''
    pf = p.paragraph_format

    if 'Heading 1' in style_name:
        pf.space_before = Pt(24)
        pf.space_after = Pt(12)
    elif 'Heading 2' in style_name:
        pf.space_before = Pt(18)
        pf.space_after = Pt(8)
    elif 'Heading 3' in style_name:
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
    else:
        if p.text.strip():
            pf.space_before = Pt(0)
            pf.space_after = Pt(6)
            pf.line_spacing = Pt(18)  # 1.5 line spacing approx

# ============================================================
# 5. 頁面邊距設定
# ============================================================
print("\n5. 設定頁面邊距...")
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ============================================================
# Save
# ============================================================
print(f"\nSaving to {OUTPUT}...")
doc.save(OUTPUT)
print("Done!")

# Quick verify
doc2 = Document(OUTPUT)
red_count = 0
blue_count = 0
for p in doc2.paragraphs:
    for elem in p._element.iter():
        if elem.tag == qn('w:color'):
            val = elem.get(qn('w:val'), '').upper()
            if val == 'FF0000': red_count += 1
            if val == '0000FF': blue_count += 1
for t in doc2.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for elem in p._element.iter():
                    if elem.tag == qn('w:color'):
                        val = elem.get(qn('w:val'), '').upper()
                        if val == 'FF0000': red_count += 1
                        if val == '0000FF': blue_count += 1

print(f"\n驗證：紅色殘留={red_count}, 藍色殘留={blue_count}")
if red_count == 0 and blue_count == 0:
    print("✓ 全部已轉為黑色")
else:
    print("⚠ 仍有非黑色字體")
