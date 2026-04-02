#!/usr/bin/env python3
"""v11: 10 modifications — format unification, content refinement, table captions,
fill 4.1/4.3, literature support, reference summaries."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
import glob
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# I/O & Config
# ============================================================
input_files = [f for f in glob.glob('*v10*0331*.docx') if not f.startswith('~')]
INPUT = input_files[0]
OUTPUT = '農糧署研究計畫書_v11_0401.docx'
print(f'Input:  {INPUT}')
print(f'Output: {OUTPUT}')

BODY_PT = 12
TABLE_PT = 12
H1_PT, H2_PT, H3_PT = 16, 14, 13
EA_FONT = '標楷體'
ASCII_FONT = 'Times New Roman'
HEADER_FILL = '2E75B6'
ODD_FILL = 'DEEAF1'
EVEN_FILL = 'FFFFFF'

doc = Document(INPUT)

# ============================================================
# Helpers
# ============================================================

def find_para(doc, text_contains, start_from=0):
    for i, p in enumerate(doc.paragraphs):
        if i >= start_from and text_contains in p.text:
            return p, i
    return None, -1

def find_heading(doc, text_contains):
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') and text_contains in p.text:
            return p
    return None

def insert_after(doc, anchor_para, text, style='Normal'):
    new = doc.add_paragraph(text, style=style)
    doc.element.body.remove(new._element)
    anchor_para._element.addnext(new._element)
    return new

def insert_before_elem(doc, anchor_elem, text, style='Normal'):
    new = doc.add_paragraph(text, style=style)
    doc.element.body.remove(new._element)
    anchor_elem.addprevious(new._element)
    return new

def insert_after_elem(doc, anchor_elem, text, style='Normal'):
    new = doc.add_paragraph(text, style=style)
    doc.element.body.remove(new._element)
    anchor_elem.addnext(new._element)
    return new

def replace_in_para(para, old, new):
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    runs = para.runs
    if not runs:
        return False
    runs[0].text = new_full
    for r in runs[1:]:
        r.text = ''
    return True

def set_east_asia_font(run, font_name):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def set_cell_shading(cell, fill_color):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)

def create_table_after(doc, anchor_elem, headers, rows):
    """Create a table and insert it after anchor_elem in the body."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        cell.paragraphs[0].add_run(h)
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            cell.paragraphs[0].add_run(val)
    doc.element.body.remove(table._tbl)
    anchor_elem.addnext(table._tbl)
    return table

def get_body_elem_text(elem):
    """Get plain text from a body-level XML element (paragraph only)."""
    if elem.tag == qn('w:p'):
        return ''.join(t.text or '' for t in elem.iter(qn('w:t')))
    return ''

def get_table_caption(table):
    """Identify a table by its header row and return the appropriate caption."""
    if not table.rows:
        return None
    cells = [cell.text.strip() for cell in table.rows[0].cells]
    h = cells[0] if cells else ''
    h2 = cells[1] if len(cells) > 1 else ''
    all_h = ' '.join(cells)

    if '亮點' in h:
        return '表 1-1　計畫主要亮點彙整'
    if '研究目的' in h:
        return '表 1-2　研究目的與對應研究問題'
    if '方法' in h and ('類別' in h or '限制' in all_h):
        return '表 2-1　三種預測方法優劣比較'
    if '描述' in h and '本研究回應' in h2:
        return '表 2-2　研究缺口與本研究回應'
    if h == '模型' and '技術類型' in h2:
        return '表 2-3　三種預測模型規格'
    if '蔬菜' in h and '旺季' in h2:
        return '表 2-4　五種蔬菜選定理由'
    if '蔬菜' in h and '交易' in h2:
        return '表 2-5　五種蔬菜交易資料盤點'
    if '資料類別' in h and '資料筆數' in h2:
        return '表 2-6　輔助資料來源'
    if '資料類別' in h and '來源機關' in h2:
        return '表 2-7　政府開放資料來源彙整'
    if '功能面向' in h:
        return '表 3-1　現有平臺與本計畫功能比較'
    if '階段' in h and '輸入' in all_h:
        return '表 3-2　預測模型架構三階段摘要'
    if '工作項目' in h and '內容說明' in h2:
        return '表 4-1　工作項目與預期交付物'
    if '指標項目' in h:
        return '表 4-2　關鍵績效指標'
    if '階段' in h and '月份' in h2:
        return '表 4-3　專案時程規劃'
    return None


# ############################################################
#  PHASE 1 — Content modifications (text-level)
# ############################################################

# --- Task 2: Simplify 1.3.1 核心問題 ---
print('\n── Task 2: Simplifying 1.3.1 ──')

p, _ = find_para(doc, '臺灣蔬菜產銷長期存在價格劇烈波動')
if p:
    replace_in_para(p, p.text, '臺灣蔬菜產銷面臨以下三項核心問題：')
    print('  ✓ 導言已精簡')

p, _ = find_para(doc, '問題一：農民缺乏未來價格趨勢')
if p:
    replace_in_para(p, p.text,
        '問題一：農民缺乏價格預測工具，無法於種植決策前掌握市場趨勢。')
    print('  ✓ 問題一已精簡')

p, _ = find_para(doc, '問題二：現有平臺僅提供歷史查詢')
if p:
    replace_in_para(p, p.text,
        '問題二：現有平臺僅供歷史查詢，缺乏整合氣象與颱風因素之前瞻預測功能。')
    print('  ✓ 問題二已精簡')

p, _ = find_para(doc, '問題三：學術研究之預測模型多停留')
if p:
    replace_in_para(p, p.text,
        '問題三：學術預測模型未能轉化為可操作之決策支援工具。')
    print('  ✓ 問題三已精簡')


# --- Task 3: Modify RQ1 sub-questions ---
print('\n── Task 3: Modifying RQ1 sub-questions ──')

p_1a, _ = find_para(doc, '子問題 1a')
if p_1a:
    replace_in_para(p_1a, p_1a.text,
        '子問題 1a：如何建立不同來源資料表之間的關聯機制'
        '（如交易行情與氣象觀測之時間—空間對應），'
        '以形成可供模型訓練之整合特徵矩陣？')
    print('  ✓ 1a 已改為強調資料表關聯')

p_1b, _ = find_para(doc, '子問題 1b')
if p_1b:
    replace_in_para(p_1b, p_1b.text,
        '子問題 1b：不同資料源之時間粒度（日／旬／月／年）'
        '與更新頻率差異如何對齊，以產出一致之多尺度分析資料？')
    new_1c = insert_after(doc, p_1b,
        '子問題 1c：氣象觀測站與行政區域之空間對應關係如何建立？')
    print('  ✓ 1b 已修改，1c 已新增')


# --- Task 4: Modify RQ2 sub-questions ---
print('\n── Task 4: Modifying RQ2 sub-questions ──')

p_2a, _ = find_para(doc, '子問題 2a')
if p_2a:
    replace_in_para(p_2a, p_2a.text,
        '子問題 2a：在相同特徵矩陣與訓練／測試分割條件下，'
        '三種模型之預測精度（MAPE、RMSE、R²）是否存在顯著差異？')
    print('  ✓ 2a 已加入評估指標')

p_2b, _ = find_para(doc, '子問題 2b')
if p_2b:
    insert_after(doc, p_2b,
        '子問題 2c：在不同時間尺度（日、週、月）與不同蔬菜品項之交叉組合下，'
        '各模型之最適應用情境為何？如何透過集成策略進一步提升整體預測表現？')
    print('  ✓ 2c 已新增（模型比較與集成策略）')


# --- Task 7: Add literature to 3.2 ---
print('\n── Task 7: Adding citations to 3.2 ──')

p, _ = find_para(doc, '本計畫採用三層式系統架構設計')
if p:
    replace_in_para(p, p.text,
        '本計畫採用三層式系統架構設計，參考現代農業決策支援系統之分層設計原則'
        '（Zhai et al., 2020），將資料蒐集、模型運算與使用者介面明確分離，'
        '以確保各層可獨立開發、測試與部署。此架構與 Rupnik et al.（2019）'
        '提出之 AgroDSS 雲端決策支援工具箱相似，均採用資料管理、分析運算'
        '與使用者介面三層分離之設計。')
    print('  ✓ 架構概述：已加入 Zhai (2020) & Rupnik (2019)')

p, _ = find_para(doc, '運算層為系統之核心')
if p:
    old = p.text
    replace_in_para(p, old,
        old.rstrip('。') + '。此設計參考 Chen et al.（2021）'
        '所建構之自動化農產品價格預測系統，該系統同樣比較多種機器學習方法'
        '（ARIMA、SVR、Prophet、XGBoost、LSTM）並以 Web 介面呈現預測結果。')
    print('  ✓ 運算層：已加入 Chen (2021)')


# --- Task 10: Reference summaries + new references ---
print('\n── Task 10: Reference summaries ──')

ref_summaries = [
    ('今周刊', '→ 探討臺灣菜價波動原因與供苗預警機制之不足。'),
    ('高雄市政府農業局', '→ 高雄市農業局之農業資訊查詢平臺。'),
    ('雲林縣政府', '→ 雲林縣政府之數位農業行動服務平臺。'),
    ('焦鈞', '→ 分析高麗菜產銷失衡之結構性原因與解方。'),
    ('報導者', '→ 調查高麗菜價崩危機與源頭管理問題。'),
    ('農傳媒', '→ 報導高麗菜菜金菜土循環與農民困境。'),
    ('臺中市政府', '→ 臺中市政府農業資料整合查詢平臺。'),
    ('Chen, T., & Guestrin', '→ 提出 XGBoost 梯度提升樹演算法。'),
    ('Hsieh, C.-Y.', '→ 以 Sigmoid 生長曲線建立嫁接適期模型。'),
    ('Hwang, Y.-W.', '→ 研究氣候變遷下開花溫度對荔枝產量之影響。'),
    ('Hyndman, R. J.', '→ 系統性比較預測精度指標之特性與適用性。'),
    ('Peng, Y.-H.', '→ 臺灣最早以多種演算法比較農產品價格預測之研究。'),
    ('Su, Y.-J.', '→ 實證分析颱風對甘藍批發市場之需求面衝擊。'),
    ('Taylor, S. J.', '→ 提出 Prophet 可分解式時間序列預測模型。'),
    ('Rumelhart, D. E.', '→ 提出反向傳播演算法，奠定神經網路訓練基礎。'),
    ('Pedregosa, F.', '→ 發表 Scikit-learn 機器學習套件。'),
    ('Sun, F.', '→ 系統性回顧農產品價格預測方法之發展趨勢。'),
    ('Paul, R. K.', '→ 以印度茄子市場驗證機器學習於蔬菜價格預測之適用性。'),
    ('Zhao, T.', '→ 提出 TCN-XGBoost 混合模型用於農產品價格預測。'),
    ('Li, B.', '→ 以最佳化神經網路組合模型預測蔬菜價格。'),
    ('Prathilothamai, M.', '→ 以 Prophet 等時間序列模型預測印度番茄市場價格。'),
]

# Find reference section (use find_para — identity check on Paragraph wrappers is unreliable)
_, ref_start_idx = find_para(doc, '參考文獻')

summary_count = 0
last_ref_para = None
if ref_start_idx >= 0:
    print(f'  Found 參考文獻 heading at paragraph index {ref_start_idx}')
    for i in range(ref_start_idx + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        text = p.text.strip()
        if not text:
            continue
        last_ref_para = p
        for key, summary in ref_summaries:
            if key in text:
                br_run = p.add_run()
                br_run.add_break()
                p.add_run('　　' + summary)
                summary_count += 1
                break
    print(f'  ✓ Added {summary_count} summaries to existing references')

# Add 3 new references
new_refs = [
    ('Chen, Z., Goh, H. S., Sin, K. L., Lim, K., Chung, N. K. H., & Liew, X. Y. '
     '(2021). Automated agriculture commodity price prediction system with machine '
     'learning techniques. Advances in Science, Technology and Engineering Systems '
     'Journal, 6(4), 376–384. https://doi.org/10.25046/aj060442',
     '→ 比較五種 ML 方法建構自動化農產品價格預測系統。'),
    ('Rupnik, R., Kukar, M., Vracar, P., Kosir, D., Pevec, D., & Bosnic, Z. '
     '(2019). AgroDSS: A decision support system for agriculture and farming. '
     'Computers and Electronics in Agriculture, 161, 260–271. '
     'https://doi.org/10.1016/j.compag.2018.04.001',
     '→ 提出 AgroDSS 雲端農業決策支援工具箱。'),
    ('Zhai, Z., Martínez, J. F., Beltran, V., & Martínez, N. L. (2020). '
     'Decision support systems for agriculture 4.0: Survey and challenges. '
     'Computers and Electronics in Agriculture, 170, 105256. '
     'https://doi.org/10.1016/j.compag.2020.105256',
     '→ 調查農業 4.0 決策支援系統之架構模式與挑戰。'),
]

if last_ref_para:
    anchor = last_ref_para
    for citation, summary in new_refs:
        new_p = insert_after(doc, anchor, citation)
        br_run = new_p.add_run()
        br_run.add_break()
        new_p.add_run('　　' + summary)
        anchor = new_p
    print(f'  ✓ Added 3 new references (Chen 2021, Rupnik 2019, Zhai 2020)')


# --- Fix existing table number references ---
print('\n── Fixing table number references ──')
fix_count = 0
for p in doc.paragraphs:
    if '表 1-1' in p.text:
        replace_in_para(p, '表 1-1', '表 3-1')
        fix_count += 1
    if '表 1-2' in p.text:
        replace_in_para(p, '表 1-2', '表 2-7')
        fix_count += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for cp in cell.paragraphs:
                if '表 1-1' in cp.text:
                    replace_in_para(cp, '表 1-1', '表 3-1')
                    fix_count += 1
                if '表 1-2' in cp.text:
                    replace_in_para(cp, '表 1-2', '表 2-7')
                    fix_count += 1
print(f'  ✓ Fixed {fix_count} table references')


# ############################################################
#  PHASE 2 — Structural changes
# ############################################################

# --- Task 5: Swap tables in section 2.5 ---
print('\n── Task 5: Swapping tables in 2.5 ──')

body = doc.element.body
all_elems = list(body)

idx_inventory = None   # "五種蔬菜之資料盤點"
idx_rationale = None   # "五種蔬菜之選定理由"
idx_auxiliary = None    # "輔助資料來源" (short title)

for i, elem in enumerate(all_elems):
    text = get_body_elem_text(elem).strip()
    if '五種蔬菜之資料盤點' in text and idx_inventory is None:
        idx_inventory = i
    elif '五種蔬菜之選定理由' in text and idx_rationale is None:
        idx_rationale = i
    elif text == '輔助資料來源' and idx_auxiliary is None:
        idx_auxiliary = i

if idx_inventory and idx_rationale and idx_auxiliary:
    block_a = all_elems[idx_inventory:idx_rationale]   # data inventory block
    block_b = all_elems[idx_rationale:idx_auxiliary]    # selection rationale block

    for elem in block_a + block_b:
        body.remove(elem)

    anchor = all_elems[idx_inventory - 1]
    # Insert block_b first (rationale), then block_a (inventory) → rationale before inventory
    for elem in reversed(block_a):
        anchor.addnext(elem)
    for elem in reversed(block_b):
        anchor.addnext(elem)

    print(f'  ✓ Swapped: rationale ({len(block_b)} elems) now before inventory ({len(block_a)} elems)')
else:
    print(f'  ⚠ Markers not found: inv={idx_inventory}, rat={idx_rationale}, aux={idx_auxiliary}')


# --- Task 8: Add summary table to 3.3 ---
print('\n── Task 8: Clarifying 3.3 with summary table ──')

p, _ = find_para(doc, '本計畫之預測模型架構分為特徵工程')
if p:
    replace_in_para(p, p.text,
        '本計畫之預測模型架構分為特徵工程、模型訓練與集成預測三個階段，'
        '各階段之工作內容摘要如下表，詳細說明見後續段落。')
    create_table_after(doc, p._element,
        ['階段', '主要工作', '輸入', '輸出'],
        [
            ['特徵工程',
             '萃取 5 類特徵：滯後值、滾動統計量、日曆、氣象、颱風',
             '原始交易與氣象資料', '特徵矩陣'],
            ['模型訓練',
             'Prophet／XGBoost／ANN 各自獨立訓練',
             '特徵矩陣', '三組預測值'],
            ['集成預測',
             'MAPE 倒數加權平均＋信賴區間',
             '三組預測值', '最終預測值與區間'],
        ])
    print('  ✓ Summary table inserted after 3.3 intro')


# --- Task 9: Fill 4.1 and 4.3 ---
print('\n── Task 9: Filling 4.1 and 4.3 ──')

p_41, _ = find_para(doc, '本計畫之主要工作項目如下')
if p_41:
    create_table_after(doc, p_41._element,
        ['工作項目', '內容說明', '預期交付物'],
        [
            ['W1：多源資料蒐集與整合',
             '整合 AMIS 交易行情、CWA 氣象觀測、農情產量統計及颱風歷史資料，'
             '建立自動化 ETL 管道與結構化資料庫',
             '結構化資料庫、資料品質檢核報告'],
            ['W2：特徵工程與資料前處理',
             '設計滯後特徵、滾動統計量、日曆特徵、氣象特徵及七維颱風特徵向量，'
             '產出標準化特徵矩陣',
             '特徵工程模組、特徵欄位說明文件'],
            ['W3：預測模型開發與評估',
             '實作 Prophet、XGBoost、ANN 三種模型，針對五種蔬菜進行三種時間尺度預測，'
             '系統性比較模型表現',
             '模型程式碼、績效比較分析報告'],
            ['W4：決策支援平臺開發',
             '開發 Web 系統，含互動式地圖、預測圖表、颱風情境模擬、預警燈號',
             '前後端程式碼、API 文件、使用手冊'],
            ['W5：系統測試與結案',
             '功能測試、整合測試、使用者驗收測試，撰寫結案報告',
             '測試報告、結案報告'],
        ])
    print('  ✓ 4.1 work items table inserted')

p_43, _ = find_para(doc, '本計畫之關鍵績效指標如下')
if p_43:
    create_table_after(doc, p_43._element,
        ['指標項目', '目標值', '量測方式'],
        [
            ['預測精度',
             '五種蔬菜平均 MAPE < 15%',
             '以擴展窗口交叉驗證計算各模型之 MAPE、RMSE、R²'],
            ['資料涵蓋率',
             '5 蔬菜 × 3 時間尺度全部完成',
             '確認 15 組預測組合均產出有效結果'],
            ['模型比較完整性',
             '完成三模型系統性比較分析',
             '產出各模型在不同品項與尺度之績效比較矩陣'],
            ['系統可用性',
             '決策支援平臺可正常運行',
             '前後端功能測試通過率 100%'],
            ['文件交付',
             '結案報告與技術文件齊備',
             '依計畫要求繳交所有交付物'],
        ])
    print('  ✓ 4.3 KPI table inserted')


# ############################################################
#  PHASE 3 — Global formatting (run LAST)
# ############################################################

# --- Task 6: Add table captions ---
print('\n── Task 6: Adding table captions ──')

tables_to_caption = []
for table in doc.tables:
    cap = get_table_caption(table)
    if cap:
        tables_to_caption.append((table, cap))

# Captions go BELOW the table (user requirement)
# Process from last to first to avoid index shifting
for table, caption in reversed(tables_to_caption):
    tbl_elem = table._tbl

    # If there's an existing caption-like paragraph ABOVE the table, remove it
    prev = tbl_elem.getprevious()
    if prev is not None and prev.tag == qn('w:p'):
        prev_text = get_body_elem_text(prev).strip()
        if prev_text.startswith('表 ') and '　' in prev_text:
            doc.element.body.remove(prev)
            print(f'  ✗ Removed old above-caption for {caption}')

    # Insert caption BELOW the table
    new_p = insert_after_elem(doc, tbl_elem, caption)
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(f'  ✓ {caption} (inserted below table)')


# --- Task 1: Unify all formatting ---
print('\n── Task 1: Unifying formatting ──')

# 1a. Apply blue theme to all tables
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            # Cell shading
            if ri == 0:
                set_cell_shading(cell, HEADER_FILL)
            elif ri % 2 == 1:
                set_cell_shading(cell, ODD_FILL)
            else:
                set_cell_shading(cell, EVEN_FILL)

            # Cell text font
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(TABLE_PT)
                    r.font.name = ASCII_FONT
                    set_east_asia_font(r, EA_FONT)
                    if ri == 0:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    else:
                        r.bold = False
                        r.font.color.rgb = RGBColor(0, 0, 0)

print(f'  ✓ Blue theme applied to {len(doc.tables)} tables')

# 1b. Set paragraph fonts
for p in doc.paragraphs:
    style = p.style.name if p.style else ''
    for r in p.runs:
        if 'Heading 1' in style:
            r.font.size = Pt(H1_PT)
        elif 'Heading 2' in style:
            r.font.size = Pt(H2_PT)
        elif 'Heading 3' in style:
            r.font.size = Pt(H3_PT)
        else:
            r.font.size = Pt(BODY_PT)
        r.font.name = ASCII_FONT
        set_east_asia_font(r, EA_FONT)
        r.font.color.rgb = RGBColor(0, 0, 0)

print('  ✓ Paragraph fonts unified')


# ============================================================
# Save
# ============================================================
doc.save(OUTPUT)
print(f'\n✅ Saved: {OUTPUT}')


# ============================================================
# Verification
# ============================================================
print('\n' + '=' * 60)
print('VERIFICATION')
print('=' * 60)

doc2 = Document(OUTPUT)

# 1. Structure
print('\n── Document Structure ──')
for p in doc2.paragraphs:
    if p.style.name.startswith('Heading'):
        level = p.style.name.replace('Heading ', '').strip()
        indent = '  ' * (int(level) - 1) if level.isdigit() else ''
        print(f'{indent}{p.text}')

# 2. Table count and captions
print(f'\n── Tables: {len(doc2.tables)} ──')
for i, t in enumerate(doc2.tables):
    cap = get_table_caption(t)
    first_cell = t.rows[0].cells[0].text[:30] if t.rows else '?'
    print(f'  [{i:2d}] {cap or "UNKNOWN"} | header[0]: {first_cell}')

# 3. Blue theme check
print('\n── Blue theme check ──')
all_ok = True
for i, t in enumerate(doc2.tables):
    if t.rows:
        tc = t.rows[0].cells[0]._tc
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            shd = tcPr.find(qn('w:shd'))
            if shd is not None:
                fill = shd.get(qn('w:fill'))
                if fill != HEADER_FILL:
                    print(f'  ⚠ Table {i}: header fill={fill}')
                    all_ok = False
if all_ok:
    print('  ✓ All tables have blue header (#2E75B6)')

# 4. Reference summaries
print('\n── Reference summaries ──')
arrow_count = sum(1 for p in doc2.paragraphs if '→' in p.text)
print(f'  Arrow summaries found: {arrow_count}')

# 5. New sub-questions check
print('\n── New sub-questions ──')
for label in ['子問題 1c', '子問題 2c']:
    p, _ = find_para(doc2, label)
    print(f'  {label}: {"✓ found" if p else "⚠ NOT FOUND"}')

# 6. 4.1 / 4.3 tables
print('\n── 4.1 / 4.3 content ──')
for label in ['W1：多源資料', '預測精度']:
    found = False
    for t in doc2.tables:
        for row in t.rows:
            for cell in row.cells:
                if label in cell.text:
                    found = True
    print(f'  {label}: {"✓ found in table" if found else "⚠ NOT FOUND"}')

print('\n✅ Verification complete')
