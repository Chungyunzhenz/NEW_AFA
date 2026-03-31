"""
修改 農糧署研究計畫書_v3_0329.docx
所有修改內容以紅色字體標示
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xFF, 0x00, 0x00)
INPUT_FILE = '農糧署研究計畫書_v3_0329.docx'
OUTPUT_FILE = '農糧署研究計畫書_v4_0330.docx'

# ============================================================
# Helper Functions
# ============================================================

def make_run_element(text, bold=False, color='FF0000', font_size=None):
    """Create a w:r element with specified formatting."""
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    if font_size:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size * 2))  # half-points
        rPr.append(sz)
    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    return run


def insert_para_after(ref_elem, text, style_id=None, bold=False):
    """Insert a paragraph after ref_elem (can be paragraph._element or any XML element)."""
    new_p = OxmlElement('w:p')
    if style_id:
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_id)
        pPr.append(pStyle)
        new_p.append(pPr)
    if text:
        new_p.append(make_run_element(text, bold=bold))
    ref_elem.addnext(new_p)
    return new_p


def insert_empty_para(ref_elem):
    """Insert an empty paragraph after ref_elem."""
    new_p = OxmlElement('w:p')
    ref_elem.addnext(new_p)
    return new_p


def insert_table_after(ref_elem, headers, rows):
    """Insert a table with red text after ref_elem."""
    tbl = OxmlElement('w:tbl')

    # Table properties
    tblPr = OxmlElement('w:tblPr')
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.append(tblStyle)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    # Table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    tbl.append(tblPr)

    # Table grid
    tblGrid = OxmlElement('w:tblGrid')
    for _ in headers:
        gridCol = OxmlElement('w:gridCol')
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    def make_cell(text, is_header=False):
        tc = OxmlElement('w:tc')
        p = OxmlElement('w:p')
        p.append(make_run_element(str(text), bold=is_header))
        tc.append(p)
        return tc

    # Header row
    tr = OxmlElement('w:tr')
    for h in headers:
        tr.append(make_cell(h, is_header=True))
    tbl.append(tr)

    # Data rows
    for row_data in rows:
        tr = OxmlElement('w:tr')
        for cell_text in row_data:
            tr.append(make_cell(cell_text))
        tbl.append(tr)

    ref_elem.addnext(tbl)
    return tbl


def set_para_red(para, new_text, bold=False):
    """Replace all runs in a paragraph with red text."""
    p_elem = para._element
    # Remove all existing runs
    for child in list(p_elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            p_elem.remove(child)
    p_elem.append(make_run_element(new_text, bold=bold))


def set_cell_text_red(cell, text, bold=False):
    """Set cell text to red."""
    for p in cell.paragraphs:
        for child in list(p._element):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'r':
                p._element.remove(child)
    cell.paragraphs[0]._element.append(make_run_element(text, bold=bold))


# ============================================================
# Load Document
# ============================================================
print("Loading document...")
doc = Document(INPUT_FILE)
paras = doc.paragraphs
tables = doc.tables

# ============================================================
# PART 1: Modify existing paragraphs (Chapters 1 & 2)
# ============================================================
print("Part 1: Modifying existing paragraphs...")

# --- Paragraph [26]: Replace model description ---
p26_new = (
    "綜合觀之，目前臺灣學術文獻中尚無研究同時達成以下三項條件：採用多模型機器學習方法、"
    "整合氣象與產銷多源開放資料、針對甘藍進行價格預測。此即本計畫之學術切入點。"
    "本研究採用 Prophet、XGBoost 與 ANN（人工神經網路）三種獨立預測模型，"
    "分別代表可分解式時間序列分析（Prophet）、梯度提升樹（XGBoost）與類神經網路（ANN）三種不同技術路線，"
    "透過多模型比較分析，探討各方法在不同蔬菜品項與市場情境下之預測表現差異。"
    "在特徵工程方面，本系統設計涵蓋時間滯後特徵、滾動統計量、日曆特徵、氣象特徵及多維度颱風特徵，"
    "為三種模型提供一致之輸入特徵矩陣，確保比較基準之公平性。"
)
set_para_red(paras[26], p26_new)
print("  [26] Updated model description")

# --- Paragraph [38]: "四種" → "三種" ---
p38_new = (
    "本節就本研究所採用之三種核心預測方法，歸納其理論依據、技術特性與農業預測情境下之適用性，"
    "作為後續第四章研究設計之理論基礎。"
)
set_para_red(paras[38], p38_new)
print("  [38] Updated method count")

# --- Paragraph [43]: Remove LightGBM, keep XGBoost ---
p43_new = (
    "隨著機器學習技術的發展，梯度提升樹（Gradient Boosting）族系模型在農產品價格預測中展現出良好的效果。"
    "Chen 與 Guestrin（2016）提出的 XGBoost，透過正則化目標函數與高效的分裂演算法，"
    "在各類預測競賽中取得優異成績。XGBoost 之核心優勢在於特徵工程的靈活性——"
    "研究者可將滯後值（lag features）、滾動統計量、日曆特徵、氣象變數等多種資訊納入模型，"
    "學習複雜的非線性交互效應，並透過特徵重要性排名提供高度可解釋之預測結果。"
)
set_para_red(paras[43], p43_new)
print("  [43] Updated ML methods (removed LightGBM)")

# --- Paragraph [44]: Update ML limitation ---
p44_new = (
    "然而，純梯度提升樹方法對時間序列的趨勢成分與季節性成分的建模不如專門的時間序列模型直觀，"
    "因此需搭配 Prophet 等可分解式模型互補。此外，樹模型之決策邊界為分段常數函數，"
    "對平滑非線性關係之捕獲不如神經網路方法，此為引入 ANN 之動機之一。"
)
set_para_red(paras[44], p44_new)
print("  [44] Updated ML limitation")

# --- Paragraph [48]: heading "集成學習方法" → "人工神經網路方法" ---
set_para_red(paras[48], "2.1.4　人工神經網路方法")
print("  [48] Changed heading to ANN")

# --- Paragraph [49]: Replace ensemble content with ANN content ---
p49_new = (
    "人工神經網路（Artificial Neural Network, ANN）為模擬生物神經元運算機制之機器學習方法，"
    "透過多層感知器（Multi-Layer Perceptron, MLP）架構，可學習輸入特徵與目標變數之間的複雜非線性映射關係。"
    "Scikit-learn 套件提供之 MLPRegressor 實作了前饋式多層感知器回歸模型，"
    "以反向傳播演算法（Backpropagation）進行權重更新，適用於中小規模資料集之回歸預測任務。"
    "相較於深度學習模型（如 LSTM、Transformer），MLP 之參數量適中，"
    "在本研究約 170 個月度資料點之規模下，較不易發生過擬合問題。"
    "ANN 與 XGBoost 之關鍵差異在於：XGBoost 產生分段常數之決策函數，"
    "而 ANN 透過連續可微之激活函數產生平滑之函數逼近，兩者在歸納偏差（inductive bias）上具有互補性。"
)
set_para_red(paras[49], p49_new)
print("  [49] Replaced ensemble with ANN content")

# --- Paragraph [51]: heading "四類" → "三種" ---
set_para_red(paras[51], "2.1.5　本研究三種預測方法綜合比較")
print("  [51] Changed comparison heading")

# --- Paragraph [84]: "多模型集成" → "多模型獨立" ---
set_para_red(paras[84], "4.4　多模型獨立預測框架")
print("  [84] Changed framework heading")

# --- Paragraph [89]: "集成權重計算" → "模型比較評估" ---
set_para_red(paras[89], "4.4.3　模型比較與評估指標")
print("  [89] Changed evaluation heading")


# ============================================================
# PART 2: Modify existing tables
# ============================================================
print("\nPart 2: Modifying existing tables...")

# --- Table 0: 功能比較表 ---
# Row 1, Col 5: "AI 多模型集成預測" → "AI 多模型獨立預測"
set_cell_text_red(tables[0].cell(1, 5), "AI 多模型獨立預測")
# Row 3, Col 5: "Prophet＋SARIMA＋XGBoost＋LightGBM 集成" → "Prophet＋XGBoost＋ANN 三模型"
set_cell_text_red(tables[0].cell(3, 5), "Prophet＋XGBoost＋ANN\n三模型獨立預測")
print("  Table 0: Updated rows 1 & 3")

# --- Table 2: 方法比較 (5 rows x 4 cols) → Rebuild ---
# Row 0: headers (keep)
# Row 1: 傳統統計 (keep as background)
# Row 2: 機器學習 → modify (remove LightGBM, add note)
set_cell_text_red(tables[2].cell(2, 0), "梯度提升樹")
set_cell_text_red(tables[2].cell(2, 1), "XGBoost\n(Chen & Guestrin, 2016)\nPython: xgboost.XGBRegressor")
set_cell_text_red(tables[2].cell(2, 2), "特徵工程靈活、可納入多種外生變數、非線性建模能力強、特徵重要性排名提供高可解釋性")
set_cell_text_red(tables[2].cell(2, 3), "決策邊界為分段常數，對平滑非線性關係捕獲不如神經網路")
# Row 3: 可分解式 (keep as is)
# Row 4: 集成學習 → 人工神經網路
set_cell_text_red(tables[2].cell(4, 0), "人工神經網路")
set_cell_text_red(tables[2].cell(4, 1), "ANN / MLP\n(Rumelhart et al., 1986)\nPython: sklearn.neural_network.MLPRegressor")
set_cell_text_red(tables[2].cell(4, 2), "學習複雜非線性映射、連續可微函數逼近、與梯度提升樹互補、中小規模資料適用")
set_cell_text_red(tables[2].cell(4, 3), "可解釋性較低、需特徵標準化前處理、超參數調校敏感")
print("  Table 2: Updated to 3 methods")

# --- Table 3: 研究缺口 ---
# Row 1: Update "四模型集成" → "三模型獨立"
set_cell_text_red(tables[3].cell(1, 1),
    "建構 Prophet、XGBoost 與 ANN 三種獨立預測模型，分別代表可分解式、梯度提升樹與神經網路三種技術路線，透過多模型比較分析驗證各方法之適用性")
print("  Table 3: Updated row 1")


# ============================================================
# PART 3: Add Task 1 tables (after Chapter 1 data section)
# ============================================================
print("\nPart 3: Adding data inventory tables...")

# Insert after paragraph [34] (which is an empty line before Chapter 2)
# Actually insert before Chapter 2, after the last content of Chapter 1
# The last content paragraph before Ch2 is around [34]-[35]
# Let's insert after paragraph [27] (the five vegetables description)

ref = paras[27]._element  # After 5 vegetables description

# Add empty line
p_empty = insert_empty_para(ref)

# Add section title
p_title = insert_para_after(p_empty, "五種蔬菜之資料盤點", bold=True)

# Add description
p_desc = insert_para_after(p_title,
    "為確認五種目標蔬菜之資料是否足以支撐預測模型訓練，以下彙整各蔬菜之資料量與完整性分析。"
    "五種蔬菜均具備 14 年以上之交易歷史資料與 16 年以上之產量統計，月度資料點約 170 個，"
    "遠超模型最低需求（24 個月），且無重大資料斷層，分析可行性高。"
)

# Data inventory table
tbl1 = insert_table_after(p_desc,
    headers=["蔬菜", "交易資料筆數", "交易時間範圍", "月度資料點", "產量資料筆數", "產量時間範圍", "資料完整性"],
    rows=[
        ["甘藍", "178,368", "2012–2026", "~170", "343", "2008–2024", "完整，無斷層"],
        ["萵苣", "268,339", "2012–2026", "~170", "326", "2008–2024", "完整，無斷層"],
        ["小白菜", "119,543", "2012–2026", "~170", "279", "2008–2024", "完整，無斷層"],
        ["花椰菜", "69,857", "2012–2026", "~170", "276", "2008–2024", "完整，無斷層"],
        ["青蔥", "138,284", "2012–2026", "~170", "334", "2008–2024", "完整，無斷層"],
    ]
)

# Table caption
p_cap1 = insert_para_after(tbl1, "資料來源：本研究整理（2026）")

# Add empty line
p_empty2 = insert_empty_para(p_cap1)

# Selection rationale table title
p_title2 = insert_para_after(p_empty2, "五種蔬菜之選定理由", bold=True)

tbl2 = insert_table_after(p_title2,
    headers=["蔬菜", "旺季月份", "淡季月份", "選定理由"],
    rows=[
        ["甘藍", "11–3 月", "6–8 月", "臺灣產量最大之葉菜類，搶種價崩案例最為典型，農糧署重點監測品項"],
        ["萵苣", "10–3 月", "6–8 月", "消費量穩定成長之葉菜類，夏季高溫減產明顯，價格季節性波動幅度大"],
        ["小白菜", "10–3 月", "6–8 月", "家庭最常購買之平價蔬菜，生長週期短（25–30 天），對短期天氣變化反應最靈敏"],
        ["花椰菜", "11–3 月", "6–9 月", "冬季主力蔬菜，近年因健康飲食趨勢需求攀升，淡季為五種最長（4 個月）"],
        ["青蔥", "10–2 月", "6–8 月", "料理不可或缺之辛香蔬菜，颱風後價格暴漲最劇烈（曾達 5–10 倍），颱風衝擊研究最佳標的"],
    ]
)
p_cap2 = insert_para_after(tbl2, "資料來源：本研究整理（2026）")

# Shared data sources
p_empty3 = insert_empty_para(p_cap2)
p_shared = insert_para_after(p_empty3,
    "此外，五種蔬菜共享之輔助資料包括：中央氣象署氣象觀測資料（119,080 筆，2005–2026 年，涵蓋溫度、降雨量、濕度等變數）、"
    "颱風歷史事件資料（145 筆，2000–2025 年，含強度等級、侵臺路徑與影響縣市）。"
    "五種蔬菜涵蓋葉菜（甘藍、萵苣、小白菜）、花菜（花椰菜）及辛香料（青蔥）三種食用類型，"
    "可全面驗證模型之跨品類適用性。"
)

print("  Added data inventory and selection rationale tables")


# ============================================================
# PART 4: Fill Chapter 3 (研究目的與問題)
# ============================================================
print("\nPart 4: Filling Chapter 3...")

# --- 3.1.1 核心問題 (after para[66]) ---
ref = paras[66]._element
p = insert_para_after(ref,
    "臺灣蔬菜產銷體系長期面臨「菜金菜土」之結構性困境——蔬菜價格在短期內大幅震盪，供需雙方均深受其害。"
    "具體而言，本研究欲解決以下三層問題："
)
p = insert_empty_para(p)
p = insert_para_after(p,
    "問題一：產銷資訊斷層。蔬菜從產地到市場的供需資訊缺乏前瞻性預測工具。"
    "農民在決定種植面積時，僅能參考「當下」或「上一期」的市場行情，無法預見未來 1 至 6 個月的市場走勢，"
    "導致「看到高價就搶種，收成時價格已崩」的惡性循環反覆發生。"
)
p = insert_empty_para(p)
p = insert_para_after(p,
    "問題二：極端天氣應變不足。臺灣每年颱風季（6 至 10 月）對蔬菜供給造成劇烈衝擊，"
    "但現有工具無法量化「不同強度的颱風，對不同蔬菜品項，在颱風後 1 至 2 個月內的價格影響幅度」。"
    "市場管理者與政策制定者在颱風來臨前缺乏科學化的情境推估依據。"
)
p = insert_empty_para(p)
p = insert_para_after(p,
    "問題三：決策工具缺乏。現有農業公開平臺（如農業部 AMIS）僅提供歷史資料查詢功能，"
    "不具備預測分析、視覺化比較與預警通知能力。學術研究中的預測模型多停留在論文階段，"
    "未能轉化為可供產業使用者直接操作的工具。"
)
print("  3.1.1 Core problems filled")

# --- 3.1.2 服務對象 (after para[68]) ---
ref = paras[68]._element
p = insert_para_after(ref,
    "本系統之設計針對以下三類使用者，各自對應不同的功能需求："
)
tbl_svc = insert_table_after(p,
    headers=["服務對象", "角色描述", "核心需求", "系統對應功能"],
    rows=[
        ["蔬菜農民", "需決定種植品項與面積之第一線生產者",
         "未來 1–3 個月，所種蔬菜價格會漲還是跌？",
         "價格預測趨勢圖、產銷預警燈號（紅黃綠）、季節性分析"],
        ["批發市場管理者", "全臺 20 個蔬果批發市場之調度管理人員",
         "颱風來臨前，哪些蔬菜可能短缺？各縣市供給量如何分佈？",
         "颱風情境模擬、臺灣縣市地圖、跨市場比較、交易量預測"],
        ["農業政策制定者", "農業部、各縣市農業局之政策規劃人員",
         "目前哪些品項有供過於求風險？是否需啟動保價收購？",
         "多品項儀表板總覽、歷史趨勢與預測對照、資料品質監控"],
    ]
)
print("  3.1.2 Service targets filled")

# --- 3.2 研究目的 (after para[70]) ---
ref = paras[70]._element
p = insert_para_after(ref,
    "本研究旨在以五種臺灣關鍵蔬菜（甘藍、萵苣、小白菜、花椰菜、青蔥）為研究對象，"
    "設計並實現一套基於多模型比較之蔬菜產銷預測與決策支援平臺。具體研究目的如下："
)
tbl_obj = insert_table_after(p,
    headers=["研究目的", "具體內容", "對應 RQ"],
    rows=[
        ["目的一：建立多源農業資料整合管道",
         "整合農糧署 AMIS 交易系統、中央氣象署 CWA 觀測資料、颱風歷史資料庫及農情產量統計等四大外部資料源，"
         "建立自動化之資料擷取、清洗與入庫流程，解決民國／西元紀年差異、日度／月度粒度不一致等異質性問題",
         "RQ1"],
        ["目的二：發展多模型獨立預測框架",
         "設計包含 Prophet、XGBoost 與 ANN 三種獨立預測模型之比較分析架構，"
         "分別代表可分解式、梯度提升樹與神經網路三種技術路線，"
         "針對五種蔬菜之月度平均價格進行 1／3／6 個月之預測，並比較各模型之預測特性與適用情境",
         "RQ2"],
        ["目的三：實現可落地之互動式決策支援平臺",
         "開發完整之 Web 應用系統，包含互動式儀表板、臺灣縣市地圖、颱風情境模擬、"
         "產銷預警燈號與資料品質監控，使預測成果能直接服務於農民、市場管理者及政策制定者",
         "RQ3"],
    ]
)
print("  3.2 Research objectives filled")

# --- 3.3 研究問題 (after para[72]) ---
ref = paras[72]._element
p = insert_para_after(ref,
    "基於上述研究目的，本研究提出以下三組研究問題："
)
p = insert_empty_para(p)
p = insert_para_after(p,
    "RQ1：如何有效整合多源異質農業資料，為五種蔬菜建立完整的產銷分析資料庫？", bold=True)
p = insert_para_after(p,
    "子問題 1a：如何處理農業部（民國紀年）與氣象署（西元紀年）之間的日期格式差異，"
    "並將日度交易、日度氣象、年度產量資料對齊至統一之月度時間尺度？")
p = insert_para_after(p,
    "子問題 1b：如何建立 22 個氣象站與 22 個縣市之空間對應關係，以支撐縣市層級的預測分析？")
p = insert_para_after(p,
    "子問題 1c：如何設計資料品質監控機制（紅黃綠燈號），幫助使用者判斷預測結果之可信度？")
p = insert_empty_para(p)
p = insert_para_after(p,
    "RQ2：三種獨立預測模型（Prophet、XGBoost、ANN）在五種蔬菜的價格預測上表現為何？", bold=True)
p = insert_para_after(p,
    "子問題 2a：三種模型各自在五種蔬菜上的預測特性與適用情境為何？是否存在「某類模型始終較優」之模式？")
p = insert_para_after(p,
    "子問題 2b：各模型是否能達到合理之預測精度目標（以 MAPE ≤ 15%、R² ≥ 0.60 為基準）？")
p = insert_para_after(p,
    "子問題 2c：多維度颱風特徵（強度等級、颱風後 1/2 個月延遲效應、極端降雨旗標）"
    "對青蔥、甘藍等颱風敏感作物之預測改善幅度為何？")
p = insert_empty_para(p)
p = insert_para_after(p,
    "RQ3：如何設計一套符合農民、市場管理者及政策制定者需求的互動式決策支援平臺？", bold=True)
p = insert_para_after(p,
    "子問題 3a：颱風情境模擬功能（可調整颱風強度與月份）如何幫助市場管理者在颱風季前進行預判與調度？")
p = insert_para_after(p,
    "子問題 3b：臺灣縣市地圖視覺化與多層次篩選功能，如何提升政策制定者之分析效率？")
p = insert_para_after(p,
    "子問題 3c：產銷預警燈號系統如何幫助農民及早識別供過於求風險？")
print("  3.3 Research questions filled")

# --- 3.4 研究範圍與限制 (after para[74]) ---
ref = paras[74]._element
p = insert_para_after(ref, "研究範圍：", bold=True)
tbl_scope = insert_table_after(p,
    headers=["維度", "範圍", "數量／規模"],
    rows=[
        ["核心研究對象", "甘藍、萵苣、小白菜、花椰菜、青蔥", "5 種蔬菜"],
        ["地理範圍", "臺灣本島及離島", "22 個縣市、20 個批發市場"],
        ["交易資料", "農糧署 AMIS 每日交易行情", "約 1,298 萬筆（2012–2026）"],
        ["氣象資料", "中央氣象署 CWA 觀測站", "約 11.9 萬筆（2005–2026）"],
        ["產量統計", "農業部年度產量資料", "約 5,951 筆（2008–2024）"],
        ["颱風事件", "中央氣象署颱風資料庫", "145 筆（2000–2025）"],
        ["預測指標", "月均價（price_avg）", "1 項（主要）"],
        ["預測期間", "短期、中期、中長期", "1／3／6 個月"],
        ["預測模型", "Prophet、XGBoost、ANN", "3 種獨立模型"],
    ]
)
p2 = insert_para_after(tbl_scope, "")
p2 = insert_para_after(p2, "研究限制：", bold=True)
tbl_limit = insert_table_after(p2,
    headers=["限制項目", "說明"],
    rows=[
        ["時間粒度", "採用月度聚合進行預測，無法捕捉日內或週內之短期價格波動"],
        ["模型選擇", "受限於月度資料點數量（約 170 個點），未採用深度學習模型（LSTM、Transformer），因其在小樣本情境下容易過擬合"],
        ["Y 變數", "以月均價格為主要預測目標，交易量預測為輔助參考"],
        ["資料庫引擎", "以 SQLite 為主，適用於開發與展示環境；大規模生產部署建議遷移至 PostgreSQL"],
    ]
)
print("  3.4 Scope and limitations filled")


# ============================================================
# PART 5: Fill Chapter 4 (研究設計)
# ============================================================
print("\nPart 5: Filling Chapter 4...")

# --- 4.1 系統架構設計 (after para[78]) ---
ref = paras[78]._element
p = insert_para_after(ref,
    "本系統採用前後端分離架構，後端負責資料管理、模型訓練與 API 服務，前端負責互動式視覺化與使用者介面。"
    "整體架構分為四層：資料來源層、資料儲存層、預測分析層與使用者介面層。"
)

# Architecture as a table (Part 1: Data tables)
p = insert_empty_para(p)
p = insert_para_after(p, "（一）資料表定義", bold=True)
p = insert_para_after(p, "本系統之核心資料表結構如下：")

tbl_db = insert_table_after(p,
    headers=["資料表", "說明", "主要欄位", "資料量"],
    rows=[
        ["trading_data\n（交易資料）", "每日批發市場交易紀錄",
         "交易日期、作物、市場、上價、中價、下價、平均價、交易量(kg)", "約 1,298 萬筆"],
        ["weather_data\n（氣象資料）", "每日氣象觀測紀錄",
         "觀測日期、縣市、平均溫度、最高溫、最低溫、降雨量(mm)、濕度(%)", "約 11.9 萬筆"],
        ["production_data\n（產量資料）", "年度農情產量統計",
         "年份、縣市、種植面積(公頃)、收穫面積(公頃)、產量(公噸)", "約 5,951 筆"],
        ["typhoon_events\n（颱風事件）", "歷史颱風紀錄",
         "颱風名稱、年份、警報起迄、強度、侵臺路徑、最大風速、最低氣壓", "145 筆"],
        ["crops（作物）", "作物主檔", "作物代碼、中文名、英文名、旺季月份、淡季月份", "22 筆"],
        ["markets（市場）", "批發市場主檔", "市場代碼、市場名稱、所在縣市", "20 筆"],
    ]
)

# Architecture Part 2: Methods
p2 = insert_para_after(tbl_db, "")
p2 = insert_para_after(p2, "（二）預測方法架構", bold=True)
p2 = insert_para_after(p2,
    "本研究採用三種獨立預測模型，各自針對相同之特徵矩陣進行訓練與預測，模型間不進行集成或加權合成。"
    "三種模型之技術定位與分工如下："
)

tbl_methods = insert_table_after(p2,
    headers=["模型", "技術類型", "Python 模組", "核心原理", "輸入資料"],
    rows=[
        ["Prophet", "可分解式時間序列", "prophet (Prophet)",
         "將價格序列分解為趨勢＋季節性＋外生回歸項三個成分，自動偵測變化點",
         "月度時間序列（ds, y）＋外生變數（天氣、颱風）"],
        ["XGBoost", "梯度提升樹", "xgboost (XGBRegressor)",
         "建立多棵決策樹逐步修正殘差，透過正則化防止過擬合，可提供特徵重要性排名",
         "特徵矩陣（20+ 維特徵：滯後值、滾動統計、日曆、天氣、颱風）"],
        ["ANN", "人工神經網路", "sklearn.neural_network\n(MLPRegressor)",
         "多層感知器架構，透過反向傳播學習特徵與目標之非線性映射關係",
         "標準化後之特徵矩陣（與 XGBoost 相同之 20+ 維特徵）"],
    ]
)

# Flow description
p3 = insert_para_after(tbl_methods,
    "預測流程：原始資料經月度聚合與特徵工程後，分別輸入三種模型進行獨立訓練。"
    "各模型各自輸出預測值與信賴區間，使用者可於儀表板上同時檢視三種模型之預測結果，"
    "比較其差異並根據特定情境選擇最適之參考依據。"
)

# Architecture Part 3: DSS Dashboard
p3 = insert_empty_para(p3)
p3 = insert_para_after(p3, "（三）DSS 決策支援儀表板", bold=True)
p3 = insert_para_after(p3,
    "本系統之前端介面為單頁式互動儀表板，主要功能區塊如下："
)
tbl_dss = insert_table_after(p3,
    headers=["功能區塊", "說明", "對應服務對象"],
    rows=[
        ["作物與時間篩選", "選擇蔬菜品項、預測期間（1/3/6 個月）、地理層級（全國/縣市/市場）", "全部"],
        ["價格趨勢預測圖", "顯示歷史實際價格與三種模型之預測曲線及信賴區間", "農民、政策制定者"],
        ["模型績效比較", "比較三種模型之 MAE、RMSE、R² 等評估指標", "政策制定者"],
        ["特徵重要性分析", "XGBoost 特徵重要性排名圖，識別影響價格之關鍵因子", "政策制定者"],
        ["臺灣縣市地圖", "以色階呈現各縣市之交易量與價格分佈", "市場管理者"],
        ["颱風情境模擬", "可調整颱風強度，模擬不同情境下之價格影響", "市場管理者"],
        ["產銷預警燈號", "紅黃綠燈號顯示供需異常警示", "農民"],
        ["資料品質監控", "紅黃綠燈號顯示各資料源之完整性與即時性", "系統管理者"],
    ]
)
print("  4.1 System architecture filled")

# --- 4.2 技術棧選型 (after para[80]) ---
ref = paras[80]._element
p = insert_para_after(ref, "本系統之技術選型如下：")
tbl_tech = insert_table_after(p,
    headers=["層級", "技術", "版本", "選用理由"],
    rows=[
        ["後端框架", "Python + FastAPI", "3.10 / 0.115", "非同步 I/O、自動 API 文件、型別驗證"],
        ["資料庫", "SQLite + SQLAlchemy", "3.45 / 2.0", "輕量部署、ORM 抽象化、可遷移至 PostgreSQL"],
        ["預測模型", "Prophet / XGBoost / MLPRegressor", "1.1 / 2.1 / sklearn 1.6", "三種互補之技術路線"],
        ["排程器", "APScheduler", "3.10", "週期性自動重訓模型"],
        ["前端框架", "React + Vite", "19.2 / 8.0", "元件化開發、熱更新"],
        ["圖表", "Recharts + D3.js", "3.8 / 7.9", "互動式圖表、臺灣地圖"],
        ["狀態管理", "Zustand", "5.0", "輕量、無樣板程式碼"],
        ["樣式", "Tailwind CSS", "4.2", "原子化 CSS、快速排版"],
    ]
)
print("  4.2 Tech stack filled")

# --- 4.3 資料庫設計 (after para[82]) ---
ref = paras[82]._element
p = insert_para_after(ref,
    "本系統之資料庫採用 SQLite 引擎，共包含 10 張資料表。核心資料表之關聯如下：")
p = insert_para_after(p,
    "（1）crops 作物主檔為核心，trading_data、production_data 透過 crop_id 外鍵關聯；"
    "（2）weather_data 透過 county_id 與 counties 縣市主檔關聯；"
    "（3）markets 市場主檔與 trading_data 透過 market_id 關聯；"
    "（4）typhoon_events 為獨立事件表，透過日期範圍與特徵工程階段納入預測模型。"
    "（5）predictions 儲存各模型之預測結果，model_registry 記錄模型訓練之績效指標與特徵重要性。"
)
print("  4.3 Database design filled")

# --- 4.4.1 框架概覽 (after para[85]) ---
ref = paras[85]._element
p = insert_para_after(ref,
    "本研究之預測框架採用「多模型獨立預測、使用者自主比較」之設計理念，"
    "以 Prophet、XGBoost 與 ANN 三種模型分別代表三種不同之技術路線：")
p = insert_para_after(p,
    "（1）Prophet：可分解式時間序列模型，擅長捕獲長期趨勢與年度季節性成分，"
    "適合季節性明顯之蔬菜品項。"
    "（2）XGBoost：梯度提升樹模型，擅長學習特徵交互效應與非線性關係，"
    "可透過特徵重要性排名提供高可解釋性。"
    "（3）ANN（MLPRegressor）：多層感知器模型，透過連續可微之激活函數進行平滑函數逼近，"
    "與 XGBoost 之分段常數決策邊界形成互補。"
)
p = insert_empty_para(p)
p = insert_para_after(p, "三種模型之選擇理由：", bold=True)
p = insert_para_after(p,
    "（1）互補性：三種模型之歸納偏差（inductive bias）各異，"
    "Prophet 擅長趨勢與季節性、XGBoost 擅長特徵交互、ANN 擅長複雜非線性映射，覆蓋不同面向。"
    "（2）方法多樣性：涵蓋可分解式、樹模型與神經網路三種技術路線，避免同質化。"
    "（3）資料規模適配：約 170 個月度資料點，深度學習（LSTM、Transformer）會過擬合，"
    "MLPRegressor 為適當之神經網路選擇。"
    "（4）可解釋性：XGBoost 提供特徵重要性、Prophet 提供成分分解、ANN 可用排列重要性（permutation importance）。"
    "（5）運算效率：不需 GPU，一般電腦即可於 5–15 分鐘內完成單一作物之全模型訓練。"
)
p = insert_empty_para(p)
p = insert_para_after(p, "不採用其他方法之理由：", bold=True)
p = insert_para_after(p,
    "（1）不採用 SARIMA：無法納入外生變數（天氣、颱風），Prophet 已涵蓋時間序列分解功能且更具彈性。"
    "（2）不採用 LightGBM：與 XGBoost 同屬梯度提升樹，功能高度重疊，無法增加方法多樣性。"
    "（3）不採用模型集成（Ensemble）：三個模型各自獨立預測，便於直接比較各模型之預測特性與適用情境，"
    "使用者可自行判斷哪個模型更適合特定情境，避免集成權重計算增加系統複雜度。"
)
print("  4.4.1 Framework overview filled")

# --- 4.4.2 特徵工程體系 (after para[87]) ---
ref = paras[87]._element
p = insert_para_after(ref,
    "三種模型共用一致之特徵工程管道，確保比較基準之公平性。特徵分為五大類：")
tbl_feat = insert_table_after(p,
    headers=["特徵類別", "特徵名稱", "說明"],
    rows=[
        ["歷史價格特徵", "lag_1, lag_2, lag_3, lag_6, lag_12", "過去 1/2/3/6/12 個月之平均價格"],
        ["滾動統計特徵", "roll_mean_3/6/12, roll_std_3/6/12", "3/6/12 個月之滾動平均值與標準差"],
        ["日曆特徵", "month, quarter, is_peak_season, sin_month, cos_month", "月份、季度、是否旺季、月份之正弦餘弦編碼"],
        ["氣象特徵", "temp_avg, rainfall_mm, temp_anomaly", "月均溫、月降雨量、溫度異常（偏離歷史均值）"],
        ["颱風特徵", "is_typhoon_month, typhoon_intensity, days_since_typhoon, post_typhoon_1m/2m, extreme_rainfall",
         "是否颱風月、颱風強度等級、距上次颱風天數、颱風後 1/2 個月旗標、極端降雨旗標"],
    ]
)
p2 = insert_para_after(tbl_feat,
    "注意：Prophet 使用時間序列格式（ds, y）搭配外生回歸變數，故僅納入氣象與颱風特徵作為額外回歸項；"
    "XGBoost 與 ANN 則使用完整之 20+ 維特徵矩陣。ANN 額外需要對特徵進行標準化（StandardScaler）前處理。"
)
print("  4.4.2 Feature engineering filled")

# --- 4.4.3 模型比較與評估指標 (after para[89]) ---
ref = paras[89]._element
p = insert_para_after(ref,
    "三種模型之預測結果將以下列評估指標進行比較：")
tbl_eval = insert_table_after(p,
    headers=["指標", "公式概念", "解讀"],
    rows=[
        ["MAE（平均絕對誤差）", "預測值與實際值之差的絕對值平均", "數值越小越好，單位與價格相同（元/公斤）"],
        ["RMSE（均方根誤差）", "預測誤差平方之均值的平方根", "對大誤差懲罰較重，單位與價格相同"],
        ["MAPE（平均絕對百分比誤差）", "百分比形式之平均誤差", "跨品項可比較，目標 ≤ 15%"],
        ["R²（決定係數）", "模型解釋之變異比例", "0–1，越接近 1 越好，目標 ≥ 0.60"],
    ]
)
p2 = insert_para_after(tbl_eval,
    "評估方式採用時間序列分割法：全部月度資料中，最後 90 天保留為完全不參與訓練之測試集，"
    "其餘資料以 70/30 比例分為訓練集與驗證集。三種模型在相同之資料分割下進行訓練與評估，確保比較之公正性。"
)
print("  4.4.3 Model evaluation filled")

# --- 4.4.4 訓練與評估流程 (after para[91]) ---
ref = paras[91]._element
p = insert_para_after(ref, "各模型之訓練配置如下：")
tbl_train = insert_table_after(p,
    headers=["模型", "主要超參數", "訓練策略"],
    rows=[
        ["Prophet", "yearly_seasonality=True, changepoint_prior_scale=0.1, 可加式季節性",
         "以完整月度序列輸入，自動偵測趨勢變化點，外生回歸項包含天氣與颱風特徵"],
        ["XGBoost", "n_estimators=200, max_depth=6, learning_rate=0.1, reg_alpha/lambda 正則化",
         "以特徵矩陣訓練，採用遞迴多步預測（recursive multi-step）策略"],
        ["ANN (MLP)", "hidden_layer_sizes=(64,32,16), activation=relu, solver=adam, early_stopping=True",
         "特徵標準化後訓練，early_stopping 防止過擬合，validation_fraction=0.15"],
    ]
)
print("  4.4.4 Training flow filled")

# --- 4.5 Task 4: 分類 vs 迴歸 & Task 5: Y 變數 (after para[93]) ---
ref = paras[93]._element
p = insert_para_after(ref,
    "本系統之模型訓練採用 APScheduler 排程器，設定每週日凌晨 02:00 自動重新訓練所有模型，"
    "確保預測結果反映最新之市場動態。此外，系統提供手動觸發 API，供管理者於需要時立即重訓。"
)

# Insert Task 4 & 5 content - add as new subsections
p = insert_empty_para(p)
p = insert_para_after(p, "預測任務類型之選擇：迴歸 vs. 分類", bold=True)
p = insert_para_after(p,
    "三種預測模型均同時支援迴歸（預測連續數值）與分類（預測類別標籤）兩種任務類型。"
    "本研究之分析如下：")
tbl_cls = insert_table_after(p,
    headers=["模型", "迴歸（Regression）", "分類（Classification）"],
    rows=[
        ["XGBoost", "XGBRegressor — 預測連續數值", "XGBClassifier — 預測類別"],
        ["ANN", "MLPRegressor — 預測連續數值", "MLPClassifier — 預測類別"],
        ["Prophet", "原生為迴歸模型", "不直接支援分類"],
    ]
)
p2 = insert_para_after(tbl_cls,
    "本研究選擇採用迴歸方向，理由如下："
    "（1）Y 變數（月平均價格）為連續數值，非離散類別；"
    "（2）迴歸能回答「價格預測為多少元/公斤」，資訊量遠大於分類之「漲或跌」；"
    "（3）若需分類功能（如價格異常警示），可由迴歸結果後處理衍生——"
    "例如預測價格偏離歷史均值超過 15% 標記為黃燈、超過 30% 標記為紅燈，"
    "此即本系統已實作之紅黃綠預警燈號機制。"
)

# Task 5: Y variable
p2 = insert_empty_para(p2)
p2 = insert_para_after(p2, "Y 變數（預測目標）之選擇", bold=True)
p2 = insert_para_after(p2,
    "本研究以「月平均價格（price_avg）」為主要預測目標變數（Y 變數），理由如下："
    "（1）價格為觸發政策介入之直接信號——價格崩跌時啟動保價收購、價格飆升時釋出庫存，"
    "對農糧署之政策決策具最高參考價值；"
    "（2）對三類服務對象（農民、市場管理者、政策制定者）均具直接之決策意義；"
    "（3）交易量（volume）雖可作為輔助參考指標，惟其政策行動對應性不如價格直接，"
    "故列為系統之輔助觀測項目，而非主要預測目標。"
)
print("  4.5 + Task 4 & 5 filled")

# --- 4.6 前端視覺化設計 (after para[95]) ---
ref = paras[95]._element
p = insert_para_after(ref,
    "前端採用 React 19 單頁式應用（SPA），以統一儀表板（Unified Dashboard）為核心頁面，"
    "整合所有分析功能於單一可捲動之介面中。主要視覺化元件包括：")
p = insert_para_after(p,
    "（1）價格趨勢預測圖（Recharts 折線圖）：顯示歷史價格與三種模型之預測曲線，附帶信賴區間色帶。"
    "（2）臺灣縣市地圖（D3.js + TopoJSON）：以色階呈現各縣市之交易量與價格分佈，支援滑鼠懸停顯示詳細數據。"
    "（3）模型績效比較圖：並列比較三種模型之 MAE、RMSE、R² 等指標。"
    "（4）特徵重要性圖（水平長條圖）：呈現 XGBoost 之特徵重要性排名，以顏色區分天氣、颱風、季節等類別。"
    "（5）季節性分析圖：逐年同月比較，識別季節性價格模式。"
    "（6）颱風情境模擬器：可調整颱風強度（輕度/中度/強烈），即時顯示模擬影響。"
)
print("  4.6 Frontend design filled")

# --- 4.7 資料品質監控 (after para[97]) ---
ref = paras[97]._element
p = insert_para_after(ref,
    "系統內建資料品質監控機制，採用紅黃綠三色燈號即時反映各資料源之健康狀態：")
tbl_dq = insert_table_after(p,
    headers=["燈號", "條件", "意義"],
    rows=[
        ["綠燈", "資料完整度 ≥ 90%，最近更新在 7 天內", "資料品質良好，預測可信度高"],
        ["黃燈", "資料完整度 70%–90%，或最近更新在 7–30 天內", "資料品質尚可，預測結果需審慎參考"],
        ["紅燈", "資料完整度 < 70%，或最近更新超過 30 天", "資料品質不足，預測結果可信度低"],
    ]
)
print("  4.7 Data quality monitoring filled")


# ============================================================
# PART 6: Fill Chapter 5 (工作內容重點)
# ============================================================
print("\nPart 6: Filling Chapter 5...")

# --- 5.1 工作項目 (after para[101]) ---
ref = paras[101]._element
p = insert_para_after(ref, "本計畫之主要工作項目如下：")
tbl_work = insert_table_after(p,
    headers=["工作項目", "內容說明", "交付物"],
    rows=[
        ["W1：多源資料整合", "整合農糧署交易、氣象署觀測、農情產量、颱風歷史等四大資料源，"
         "建立自動化 ETL 管道", "SQLite 資料庫（10 張資料表）、資料品質報告"],
        ["W2：特徵工程設計", "設計 20+ 維特徵矩陣，含歷史價格、滾動統計、日曆、氣象、颱風等五大類",
         "特徵工程模組、特徵說明文件"],
        ["W3：三模型訓練與評估", "分別實作 Prophet、XGBoost、ANN 三種模型，"
         "以 MAE、RMSE、MAPE、R² 評估預測表現", "訓練程式碼、模型績效比較報告"],
        ["W4：決策支援平臺開發", "開發 Web 應用系統，含互動式儀表板、地圖視覺化、颱風模擬、預警燈號",
         "前後端完整程式碼、API 文件"],
        ["W5：系統測試與文件撰寫", "系統功能測試、研究報告撰寫",
         "測試報告、研究計畫書（完整版）"],
    ]
)
print("  5.1 Work items filled")

# --- 5.2 預期貢獻 (after para[103]) ---
ref = paras[103]._element
p = insert_para_after(ref, "本計畫預期在學術與實務兩方面產生以下貢獻：")
p = insert_para_after(p,
    "（1）學術貢獻：首次針對臺灣五種蔬菜，系統性比較三種不同技術路線（可分解式、梯度提升樹、神經網路）"
    "之預測表現，填補現有文獻中缺乏多模型比較分析之研究缺口。"
)
p = insert_para_after(p,
    "（2）方法論貢獻：設計七維颱風特徵向量，細緻建模颱風強度、延遲效應與極端降雨等多面向影響，"
    "超越現有研究僅以二元變數（有/無颱風）粗略處理之做法。"
)
p = insert_para_after(p,
    "（3）實務貢獻：開發完整之 Web 決策支援系統，將預測模型從論文層面落地為可操作之工具，"
    "直接服務於農民、市場管理者與政策制定者三類使用者。"
)
print("  5.2 Expected contributions filled")

# --- 5.3 KPI (after para[105]) ---
ref = paras[105]._element
p = insert_para_after(ref, "本計畫之關鍵績效指標如下：")
tbl_kpi = insert_table_after(p,
    headers=["KPI", "目標值", "量測方式"],
    rows=[
        ["預測精度（MAPE）", "≤ 15%", "三種模型在五種蔬菜之平均 MAPE"],
        ["預測解釋力（R²）", "≥ 0.60", "三種模型在五種蔬菜之平均 R²"],
        ["資料整合", "≥ 1,300 萬筆", "資料庫總資料量"],
        ["API 端點", "≥ 30 個", "後端 API 總數"],
        ["系統可用性", "前後端完整運行", "系統展示驗收"],
    ]
)
print("  5.3 KPI filled")

# --- 5.4 時程規劃 (after para[107]) ---
ref = paras[107]._element
p = insert_para_after(ref, "本計畫以六個月為執行期程，各階段規劃如下：")
tbl_schedule = insert_table_after(p,
    headers=["月份", "階段", "主要工作", "里程碑"],
    rows=[
        ["第 1 個月", "資料整合", "建立 ETL 管道、整合四大資料源、資料品質檢核", "資料庫建置完成"],
        ["第 2 個月", "特徵工程", "設計與實作 20+ 維特徵、資料分割策略", "特徵矩陣產出"],
        ["第 3 個月", "模型開發", "實作 Prophet、XGBoost、ANN 三種模型", "模型訓練完成"],
        ["第 4 個月", "模型評估", "五種蔬菜之預測精度評估與模型比較分析", "績效報告完成"],
        ["第 5 個月", "系統開發", "前端儀表板、地圖、颱風模擬、預警系統", "系統展示版完成"],
        ["第 6 個月", "測試與文件", "系統測試、報告撰寫、成果展示", "計畫結案"],
    ]
)
print("  5.4 Schedule filled")


# ============================================================
# PART 7: Save
# ============================================================
print(f"\nSaving to {OUTPUT_FILE}...")
doc.save(OUTPUT_FILE)
print("Done!")
