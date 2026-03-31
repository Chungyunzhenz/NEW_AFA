"""v6→v7: 改善表格呈現 + 清理空白段落"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = 'FF0000'
INPUT = '農糧署研究計畫書_v6_0330.docx'
OUTPUT = '農糧署研究計畫書_v7_0330.docx'

def mke(text, bold=False, color=RED):
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    if bold:
        rPr.append(OxmlElement('w:b'))
    run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; t.set(qn('xml:space'), 'preserve')
    run.append(t)
    return run

def sc(cell, text, bold=False, color=RED):
    for p in cell.paragraphs:
        for ch in list(p._element):
            if (ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag) == 'r':
                p._element.remove(ch)
    cell.paragraphs[0]._element.append(mke(text, bold, color))

doc = Document(INPUT)
P = doc.paragraphs
T = doc.tables

# ============================================================
# 1. 改善表格 0（功能比較表）
# ============================================================
print("1. 改善功能比較表...")

t = T[0]
# 簡化 Row 1: 價格預測
sc(t.cell(1,1), "✗ 僅歷史查詢")
sc(t.cell(1,2), "✗ 僅產值模擬")
sc(t.cell(1,3), "✗ 僅即時行情")
sc(t.cell(1,4), "✗ 僅即時行情")
sc(t.cell(1,5), "✓ AI 多模型預測\n（Prophet、XGBoost、ANN）")

# Row 2: 預警機制
sc(t.cell(2,1), "✗")
sc(t.cell(2,2), "✗")
sc(t.cell(2,3), "△ 氣象預警")
sc(t.cell(2,4), "△ 氣象告警")
sc(t.cell(2,5), "✓ 紅黃綠燈號\n（量化觸發）")

# Row 3: AI/ML
sc(t.cell(3,1), "✗")
sc(t.cell(3,2), "✗")
sc(t.cell(3,3), "✗")
sc(t.cell(3,4), "△ 影像辨識\n（非價格預測）")
sc(t.cell(3,5), "✓ 三種 ML 模型\n獨立預測比較")

# Row 4: 氣象整合
sc(t.cell(4,1), "✗")
sc(t.cell(4,2), "△ GIS")
sc(t.cell(4,3), "△ 天氣預警")
sc(t.cell(4,4), "△")
sc(t.cell(4,5), "✓ CWA API 整合\n＋颱風情境模擬")

# Row 5: 資料縱深
sc(t.cell(5,1), "歷史查詢")
sc(t.cell(5,2), "近期資料")
sc(t.cell(5,3), "近期資料")
sc(t.cell(5,4), "近期資料")
sc(t.cell(5,5), "✓ 多源整合\n（2005–2026）")

# Row 6: 可解釋性
sc(t.cell(6,1), "✗")
sc(t.cell(6,2), "✗")
sc(t.cell(6,3), "✗")
sc(t.cell(6,4), "✗")
sc(t.cell(6,5), "✓ 特徵重要性分析")

print("   表格已簡化（✓/✗/△ 標記）")

# ============================================================
# 2. 清理空白段落
# ============================================================
print("\n2. 清理空白段落...")

body = doc.element.body
removed = 0

# Collect all paragraph elements that are empty
# We need to be careful - keep at most 1 empty paragraph between sections
paras_to_remove = []
prev_empty = False

for elem in list(body):
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    if tag == 'p':
        # Check if paragraph is empty (no text content)
        text = ''.join(t.text or '' for t in elem.iter() if t.tag.endswith('}t') or t.tag == 't').strip()

        # Check if it has a heading style (keep even if empty text)
        has_heading = False
        for pPr in elem.findall(qn('w:pPr')):
            for pStyle in pPr.findall(qn('w:pStyle')):
                val = pStyle.get(qn('w:val'), '')
                if 'Heading' in val:
                    has_heading = True

        if not text and not has_heading:
            if prev_empty:
                # This is a consecutive empty paragraph - mark for removal
                paras_to_remove.append(elem)
            prev_empty = True
        else:
            prev_empty = False
    else:
        prev_empty = False  # Reset on non-paragraph elements (tables etc)

# Remove the consecutive empty paragraphs
for elem in paras_to_remove:
    body.remove(elem)
    removed += 1

print(f"   移除 {removed} 個多餘空白段落")

# ============================================================
# 3. 確認被刪除的章節標題也清掉
# ============================================================
print("\n3. 清理殘留的空標題...")

# Find and remove empty heading paragraphs that belong to deleted Chapter 4 (研究設計)
# These are headings with no content paragraphs after them
deleted_headings = [
    '4.1　系統架構設計', '4.2　技術棧選型', '4.3　資料庫設計',
    '4.4　多模型獨立預測框架', '4.4.1　框架概覽', '4.4.2　特徵工程體系',
    '4.4.3　模型比較與評估指標', '4.4.4　訓練與評估流程',
    '4.5　自動化排程設計', '4.6　前端視覺化設計', '4.7　資料品質監控機制',
    '第四章　研究設計', '3.1.2　服務對象'
]

# Re-scan after removal
for elem in list(body):
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    if tag == 'p':
        text = ''.join(t.text or '' for t in elem.iter() if t.tag.endswith('}t') or t.tag == 't').strip()
        if text in deleted_headings:
            body.remove(elem)
            removed += 1
            print(f"   移除殘留標題: {text}")

# Also remove empty headings that were cleared but kept their style
for elem in list(body):
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    if tag == 'p':
        text = ''.join(t.text or '' for t in elem.iter() if t.tag.endswith('}t') or t.tag == 't').strip()
        if not text:
            # Check if it's a heading style
            for pPr in elem.findall(qn('w:pPr')):
                for pStyle in pPr.findall(qn('w:pStyle')):
                    val = pStyle.get(qn('w:val'), '')
                    if 'Heading' in val:
                        body.remove(elem)
                        removed += 1

print(f"   總計移除 {removed} 個元素")

# ============================================================
# 4. 清理被清空的表格
# ============================================================
print("\n4. 清理空表格...")
tbl_removed = 0
for elem in list(body):
    tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
    if tag == 'tbl':
        # Check if all cells are empty
        all_empty = True
        for t_elem in elem.iter():
            if t_elem.tag.endswith('}t') or t_elem.tag == 't':
                if t_elem.text and t_elem.text.strip() and t_elem.text.strip() not in ('', '—'):
                    all_empty = False
                    break
        if all_empty:
            body.remove(elem)
            tbl_removed += 1

print(f"   移除 {tbl_removed} 個空表格")

# ============================================================
# Save
# ============================================================
print(f"\nSaving to {OUTPUT}...")
doc.save(OUTPUT)

# Final stats
doc2 = Document(OUTPUT)
print(f"段落數: {len(doc2.paragraphs)}, 表格數: {len(doc2.tables)}")
print("Done!")
