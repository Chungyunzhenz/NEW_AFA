#!/usr/bin/env python3
"""v10: font→black, table descriptions, Ch3 architecture, prediction timeframes."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

doc = Document('農糧署研究計畫書_v9_0331.docx')

# ============================================================
# Helpers
# ============================================================

def find_para(doc, text_contains, start_from=0):
    for i, p in enumerate(doc.paragraphs):
        if i >= start_from and text_contains in p.text:
            return p, i
    return None, -1

def find_heading(doc, text_contains):
    for p in doc.paragraphs:
        if p.style.name.startswith('Heading') and text_contains in p.text:
            return p
    return None

def insert_after(doc, anchor_para, text, style='Normal'):
    """Insert a new paragraph right after anchor_para."""
    new = doc.add_paragraph(text, style=style)
    doc.element.body.remove(new._element)
    anchor_para._element.addnext(new._element)
    return new

def insert_before_elem(doc, anchor_elem, text, style='Normal'):
    """Insert a new paragraph before an XML element."""
    new = doc.add_paragraph(text, style=style)
    doc.element.body.remove(new._element)
    anchor_elem.addprevious(new._element)
    return new

def replace_in_para(para, old, new):
    """Replace text across runs. Sets full text on first run, clears rest."""
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    runs = para.runs
    if not runs:
        return False
    runs[0].text = new_full
    for r in runs[1:]:
        r.text = ''
    return True


# ============================================================
# Task 4: Prediction timeframe text replacements
# ============================================================
print('Task 4: Updating prediction timeframes...')

# 1.6 研究範圍 — main timeframe change
p, _ = find_para(doc, '月度平均價格之')
if p:
    replace_in_para(p,
        '進行月度平均價格之 1、3、6 個月預測',
        '進行 1 日、5 日（週）、28 日（月）三種時間尺度之平均價格預測')
    print('  ✓ 1.6 研究範圍：時間尺度已更新')

# 1.6 研究限制
p, _ = find_para(doc, '採月度聚合進行預測')
if p:
    full = p.text
    full = full.replace(
        '採月度聚合進行預測，無法捕捉日內或週內之短期波動',
        '預測時間尺度最細為日度，尚未涵蓋盤中即時價格變化')
    full = full.replace(
        '受限於月度資料點數量，未採用深度學習模型',
        '受限於蔬菜品項之歷史資料年限，未採用深度學習模型')
    full = full.replace(
        '以月均價格為主要預測目標',
        '以平均價格為主要預測目標')
    runs = p.runs
    if runs:
        runs[0].text = full
        for r in runs[1:]:
            r.text = ''
    print('  ✓ 1.6 研究限制：三項限制已更新')

# RQ1 子問題 1a — 月度分析 → 多尺度分析
p, _ = find_para(doc, '產出一致之月度分析資料')
if p:
    replace_in_para(p, '月度分析', '多尺度分析')
    print('  ✓ RQ1 子問題 1a：月度→多尺度')

# Scan tables for remaining "1、3、6" or "月度" references
print('\n  Scanning tables for remaining timeframe references...')
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            ct = cell.text
            if '1、3、6' in ct or '月度平均' in ct:
                print(f'  ⚠ Table {ti} row {ri} col {ci}: {ct[:80]}')
                # Auto-fix table cells too
                for cp in cell.paragraphs:
                    for cr in cp.runs:
                        cr.text = cr.text.replace('1、3、6 個月', '1 日、5 日（週）、28 日（月）')
                        cr.text = cr.text.replace('月度平均', '多尺度平均')


# ============================================================
# Task 2: Expand table descriptions
# ============================================================
print('\nTask 2: Adding table descriptions...')

# --- 蔬菜資料盤點 table: add post-table summary ---
p_veg, idx = find_para(doc, '以下為五種蔬菜之資料量與完整性')
if p_veg:
    # Find the next "資料來源" paragraph after this one
    p_src, _ = find_para(doc, '資料來源：本研究整理', start_from=idx + 1)
    if p_src:
        insert_after(doc, p_src,
            '上表顯示五種蔬菜之交易資料筆數均超過 6 萬筆，時間跨度達 14 年'
            '（2012–2026），且全部品項均無資料斷層。以甘藍為例，約 17 萬筆'
            '交易紀錄涵蓋全臺 20 個批發市場，資料規模足以支撐多時間尺度預測'
            '模型之訓練、驗證與交叉比較。')
        print('  ✓ 蔬菜資料盤點表：已加表後摘要')

# --- 選定理由 table: add intro ---
p_reason, _ = find_para(doc, '五種蔬菜之選定理由')
if p_reason:
    insert_after(doc, p_reason,
        '五種蔬菜之選定兼顧農業經濟代表性與資料可得性，涵蓋大宗葉菜'
        '（甘藍）、消費穩定型（萵苣）、短週期作物（小白菜）、季節敏感型'
        '（花椰菜）及調味類（青蔥），旺淡季分布互補，可全面檢驗模型在'
        '不同價格波動模式下之預測適用性。詳見下表。')
    print('  ✓ 選定理由表：已加表前說明')

# --- 輔助資料來源 table: add intro ---
p_aux, _ = find_para(doc, '輔助資料來源')
if p_aux and len(p_aux.text.strip()) < 20:  # short title, not a long sentence
    insert_after(doc, p_aux,
        '除批發市場交易行情與農情產量統計外，本計畫另整合中央氣象署逐日氣象'
        '觀測資料（涵蓋 22 縣市之溫度、降雨量與濕度）以及 2000 年以來共 '
        '145 筆侵臺颱風歷史事件，藉由氣象與極端天候變數補充價格波動之外生'
        '解釋力。各輔助資料之規模與涵蓋範圍彙整如下表。')
    print('  ✓ 輔助資料來源表：已加表前說明')


# ============================================================
# Task 3: Ch3 new content — 3.2, 3.3, 3.4
# ============================================================
print('\nTask 3: Adding Ch3 architecture content...')

ch4 = find_heading(doc, '第四章')
if ch4:
    # Build section content: list of (style, text)
    sections = [
        ('Heading 2', '3.2　系統整體架構'),
        ('Normal',
         '本計畫採用三層式系統架構設計，將資料蒐集、模型運算與使用者介面'
         '明確分離，以確保各層可獨立開發、測試與部署。'),
        ('Normal',
         '資料層負責多源政府開放資料之自動化擷取與整合。系統透過排程機制'
         '每日自動連線農糧署 AMIS 批發市場交易系統及中央氣象署開放資料平'
         '臺，擷取最新交易行情與氣象觀測資料，並將資料清洗、格式統一後寫'
         '入結構化資料庫。農情產量統計與颱風歷史事件則透過定期批次匯入方'
         '式更新。'),
        ('Normal',
         '運算層為系統之核心，涵蓋特徵工程、多模型訓練、集成預測與模型評'
         '估四個階段。系統透過排程機制定期啟動模型重新訓練，依各作物與地'
         '理區域之最新資料自動更新預測結果。訓練完成之模型連同評估指標一'
         '併註冊至模型倉庫，確保各版本可追溯。'),
        ('Normal',
         '展示層提供互動式 Web 儀表板，整合臺灣縣市地圖、時序預測圖表、'
         '紅黃綠預警燈號及颱風情境模擬等功能模組，使農糧署人員、農民及市'
         '場管理者可依其決策需求即時查閱預測資訊。'),
        ('Normal', ''),

        ('Heading 2', '3.3　預測模型架構'),
        ('Normal',
         '本計畫之預測模型架構分為特徵工程、模型訓練與集成預測三個階段。'),
        ('Normal',
         '特徵工程階段，系統自原始交易資料中萃取多類型特徵：（1）滯後特'
         '徵——以過去 1、2、3、6、12 期之價格作為輸入；（2）滾動統計量'
         '——計算過去 3、6、12 期之移動平均值、標準差、最大值與最小值；'
         '（3）日曆特徵——包含月份、季度、週次及旺淡季旗標；（4）氣象特'
         '徵——整合各縣市之平均溫度、降雨量與異常指標；（5）颱風特徵向量'
         '——以強度等級、侵襲路徑、距觀測日天數、颱風後 1 個月及 2 個月'
         '旗標、極端降雨量與受影響縣市數等七個維度建模。'),
        ('Normal',
         '模型訓練階段，系統獨立訓練 Prophet、XGBoost 與 ANN 三種模型。'
         'Prophet 透過趨勢分解與季節性建模捕獲長期週期；XGBoost 利用梯度'
         '提升樹學習特徵間之非線性交互效應，並輸出特徵重要性排名；ANN 以'
         '多層感知器架構逼近輸入特徵與目標變數之複雜映射關係。三種模型分'
         '別針對 1 日、5 日（週）、28 日（月）三種時間尺度之平均價格進行'
         '預測。'),
        ('Normal',
         '集成預測階段，系統以驗證集上各模型之 MAPE 倒數為權重進行加權平'
         '均，產出最終預測值。信賴區間取各模型預測區間之最寬範圍，以確保'
         '涵蓋率。模型評估採用 MSE、RMSE、MAE、R² 與 MAPE 五項指標，並'
         '搭配擴展窗口時序交叉驗證以避免資料洩漏。'),
        ('Normal', ''),

        ('Heading 2', '3.4　決策支援平臺設計'),
        ('Normal',
         '決策支援平臺以使用者導向為設計原則，提供以下核心功能模組：'),
        ('Normal',
         '（1）互動式臺灣縣市地圖：以色階呈現 22 縣市之交易量與價格分布'
         '，使用者點擊各縣市可進一步檢視該區域之詳細交易行情與預測結果。'),
        ('Normal',
         '（2）預警燈號系統：依據量化閾值設定紅、黃、綠三級警示。當預測'
         '價格偏離歷史均值超過設定標準差時觸發警示，協助使用者即時掌握價'
         '格異常訊號。'),
        ('Normal',
         '（3）颱風情境模擬器：使用者可輸入颱風強度、路徑與預計影響縣市'
         '等參數，系統即時模擬颱風事件對各作物價格之潛在衝擊，輔助事前應'
         '變決策。'),
        ('Normal',
         '（4）模型可解釋性面板：呈現 XGBoost 模型之特徵重要性排名以及集'
         '成預測之各模型權重分布，使預測結果具備可解釋性與透明度。'),
        ('Normal',
         '（5）多層級分析：支援全國、縣市、批發市場三種地理粒度之切換，'
         '滿足不同層級決策者之分析需求。'),
        ('Normal',
         '（6）資料匯出功能：預測結果與歷史交易資料均可匯出為 CSV 格式，'
         '便於後續加工分析或彙報使用。'),
        ('Normal', ''),
    ]

    # Insert each paragraph before Ch4 heading (in order)
    for style, text in sections:
        insert_before_elem(doc, ch4._element, text, style=style)

    print(f'  ✓ 已新增 3.2, 3.3, 3.4 共 {len(sections)} 個段落')


# ============================================================
# Task 1: All font colors → black
# ============================================================
print('\nTask 1: Setting all font colors to black...')

count = 0
def fix_color(run_elem):
    global count
    rPr = run_elem.find(qn('w:rPr'))
    if rPr is not None:
        color = rPr.find(qn('w:color'))
        if color is not None:
            val = color.get(qn('w:val'), '')
            if val and val.upper() != '000000':
                color.set(qn('w:val'), '000000')
                for attr in [qn('w:themeColor'), qn('w:themeTint'), qn('w:themeShade')]:
                    if attr in color.attrib:
                        del color.attrib[attr]
                count += 1

# Paragraphs
for para in doc.paragraphs:
    for run in para.runs:
        fix_color(run._element)

# Table cells
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    fix_color(run._element)

print(f'  ✓ 已修正 {count} 個非黑色 run')


# ============================================================
# Save
# ============================================================
output = '農糧署研究計畫書_v10_0331.docx'
doc.save(output)
print(f'\n✅ Saved: {output}')


# ============================================================
# Verification
# ============================================================
print('\n=== Verification ===')

doc2 = Document(output)

# 1. Check structure
print('\n--- Document Structure ---')
for p in doc2.paragraphs:
    if p.style.name.startswith('Heading'):
        indent = '  ' * (int(p.style.name.replace('Heading', '').strip()) - 1)
        print(f'{indent}{p.text}')

# 2. Check remaining non-black colors
print('\n--- Non-black color check ---')
bad = 0
for p in doc2.paragraphs:
    for r in p.runs:
        rPr = r._element.find(qn('w:rPr'))
        if rPr is not None:
            c = rPr.find(qn('w:color'))
            if c is not None:
                v = c.get(qn('w:val'), '')
                if v and v.upper() != '000000':
                    bad += 1
                    if bad <= 5:
                        print(f'  ⚠ color={v} | {r.text[:50]}')
if bad == 0:
    print('  ✓ All paragraph text is black')
else:
    print(f'  ⚠ {bad} non-black runs remaining')

# 3. Check remaining old timeframe references
print('\n--- Old timeframe reference check ---')
found_old = False
for p in doc2.paragraphs:
    if '1、3、6 個月' in p.text or '月度平均價格' in p.text:
        print(f'  ⚠ {p.text[:80]}')
        found_old = True
for t in doc2.tables:
    for row in t.rows:
        for cell in row.cells:
            if '1、3、6 個月' in cell.text or '月度平均價格' in cell.text:
                print(f'  ⚠ (table) {cell.text[:80]}')
                found_old = True
if not found_old:
    print('  ✓ No old timeframe references found')
